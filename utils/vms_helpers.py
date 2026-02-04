import os
import re
from datetime import datetime

import os
import re
import json
import time
import logging
from datetime import datetime, timedelta
from typing import List, Dict, Optional, Any, Tuple
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn
from docx.oxml import parse_xml

# Load environment variables
load_dotenv()

# Define multiple credential sets - DO NOT CHANGE USERNAME AND PASSWORD
CREDENTIAL_SETS = [
    {
        'username': os.getenv('VMS_USERNAME_1'),
        'password': os.getenv('VMS_PASSWORD_1'),
        'org_key': os.getenv('VMS_ORG_KEY_1')
    },
    {
        'username': os.getenv('VMS_USERNAME_2'),
        'password': os.getenv('VMS_PASSWORD_2'), 
        'org_key': os.getenv('VMS_ORG_KEY_2')
    }
]

STATE_MAPPING = {
    'VA': {
        'full_name': 'Virginia',
        'city': 'Richmond',
        'sm_template': 'SM_Virginia.docx',
        'rtr_template': 'RTR_Virginia.docx'
    },
    'NC': {
        'full_name': 'North Carolina',
        'city': 'Raleigh', 
        'sm_template': 'SM_North_Carolina.docx',
        'rtr_template': 'RTR_North_Carolina.docx'
    },
    'GA': {
        'full_name': 'Georgia',
        'city': 'Atlanta',
        'sm_template': 'SM_Georgia.docx',
        'rtr_template': 'RTR_Georgia.docx'
    },
    'IN': {
        'full_name': 'Indiana',
        'city': 'Indianapolis',
        'sm_template': 'SM_Indiana.docx',
        'rtr_template': 'RTR_Indiana.docx'
    },
    'FL': {
        'full_name': 'Florida',
        'city': 'Jacksonville',
        'sm_template': 'SM_Florida.docx',
        'rtr_template': 'RTR_Florida.docx'
    },
    'ID': {
        'full_name': 'Idaho',
        'city': 'Boise',
        'sm_template': 'RTR_SM_Idaho.docx',
        'rtr_template': 'RTR_SM_Idaho.docx'
    },
    'IA': {
        'full_name': 'Iowa',
        'city': 'Cedar Rapids',
        'sm_template': 'SM_Iowa.docx',
        'rtr_template': 'RTR_Iowa.docx'
    },
    'DEFAULT': {
        'full_name': 'Default',
        'city': 'Default',
        'sm_template': 'SM.docx',
        'rtr_template': 'RTR.docx'
    }
}

def extract_state_from_job_id(content):
    """Extract state abbreviation from Job ID line in content with proper prioritization"""
    lines = content.split('\n')
    

    # PRIORITY 0: Check for specific agency patterns (DBHDS = Virginia)
    for i, line in enumerate(lines[:15]):  # Check first 15 lines
        if 'DBHDS' in line.upper():
            print(f"  🏛️ Detected Virginia DBHDS agency")
            return 'VA'  # DBHDS is Virginia Department of Behavioral Health
        
    # PRIORITY 1: Look for Job ID line pattern in the FIRST FEW LINES (most reliable)
    for i, line in enumerate(lines[:10]):  # Only check first 10 lines
        if line.startswith('Job ID:') or 'Job ID:' in line:
            match = re.search(r'Job ID:\s*([A-Z]{2})-\d+', line)
            if match:
                state_abbr = match.group(1)
                if state_abbr in STATE_MAPPING:
                    print(f"  Extracted state from Job ID line: {state_abbr}")
                    return state_abbr
    
    # PRIORITY 2: Look for state pattern in the TITLE (first few lines)
    for i, line in enumerate(lines[:5]):  # Check first 5 lines for title
        # Look for patterns like: FL-DOT-, NC-FAST-, VA-123, etc.
        match = re.search(r'^([A-Z]{2})-[A-Z]', line)
        if match:
            state_abbr = match.group(1)
            if state_abbr in STATE_MAPPING:
                print(f"  Extracted state from title pattern: {state_abbr}")
                return state_abbr
    
    # PRIORITY 3: Look for state abbreviations in the title or content (avoid template text)
    for line in lines:
        # Skip lines that look like they're from templates, not actual job data
        if any(template_text in line for template_text in 
               ['VectorVMS Requirement', 'Candidate Full Legal Name', 'Candidate Pay Rate']):
            continue
            
        for state_abbr in STATE_MAPPING.keys():
            if state_abbr != 'DEFAULT' and state_abbr in line:
                # Make sure it's not part of a word and is a valid state code
                if re.search(r'\b' + state_abbr + r'\b', line):
                    print(f"  Extracted state from content: {state_abbr}")
                    return state_abbr
    
    # PRIORITY 4: Look for city names in content (avoid template cities)
    for line in lines:
        # Skip template-looking lines
        if 'Candidate' in line or 'VectorVMS' in line:
            continue
            
        for state_abbr, state_info in STATE_MAPPING.items():
            if state_abbr != 'DEFAULT' and state_info['city'].lower() in line.lower():
                print(f"  Extracted state from city: {state_abbr}")
                return state_abbr
    
    # PRIORITY 5: Look for full state names in content (avoid template text)
    for line in lines:
        if 'Managed Services Provider Contract' in line:
            continue  # Skip template text
            
        for state_abbr, state_info in STATE_MAPPING.items():
            if state_abbr != 'DEFAULT' and state_info['full_name'].lower() in line.lower():
                print(f"  Extracted state from full name: {state_abbr}")
                return state_abbr
    
    print("  Could not determine state, using DEFAULT")
    return 'DEFAULT'  # Final fallback

def validate_email_addresses(email_list):
    """Validate email addresses before sending"""
    email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    valid_emails = []
    
    for email in email_list if email_list else []:
        if isinstance(email, str) and re.match(email_regex, email):
            valid_emails.append(email)
        else:
            logging.warning(f"Invalid email address: {email}")
    
    return valid_emails

def set_table_borders(table):
    """
    Set borders for all cells in a table to maintain the table format
    """
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        # Add table borders
        tblBorders = parse_xml(r'''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
            </w:tblBorders>''')
        
        if tblPr is None:
            from docx.oxml.xmlchemy import OxmlElement
            tblPr = OxmlElement('w:tblPr')
            tbl.append(tblPr)
        
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(tblBorders)
    except:
        # If border setting fails, just continue without borders
        pass

def parse_skills_table(skills_content):
    """
    Parse the skills content - handles both markdown table and bullet formats
    """
    skills_data = []
    
    if not skills_content:
        print("  No skills content provided to parse")
        return skills_data
    
    print(f"  Parsing skills content: {len(skills_content)} chars")
    
    # Check if it's a placeholder table
    if "N/A | N/A | N/A" in skills_content:
        print("  Placeholder skills table found")
        return skills_data
    
    # Check if it's regex-extracted format (SKILLS TABLE: with bullet points)
    if "SKILLS TABLE:" in skills_content or (('Required' in skills_content or 'Desired' in skills_content) and 'Years' in skills_content and '|' not in skills_content):
        print("  🔍 Detected regex-extracted skills format")
        return parse_regex_extracted_skills(skills_content)
    
    # Otherwise, try to parse as markdown table
    lines = skills_content.split('\n')
    print(f"  Total lines in skills table: {len(lines)}")
    
    # Find the data rows (skip header and separator)
    data_lines = []
    in_data_section = False
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Markdown table separator (|---|)
        if line.startswith('|---') or '--- | ---' in line or line.startswith('|--'):
            in_data_section = True
            continue
        
        # Data rows (start with | and not header)
        if line.startswith('|') and in_data_section:
            if 'Skill' not in line and '---' not in line:
                data_lines.append(line)
    
    print(f"  Found {len(data_lines)} data rows to process")
    
    # Process each data row
    for line in data_lines:
        try:
            # Split by pipe and clean up
            parts = [part.strip() for part in line.split('|')]
            # Remove empty first and last parts (from leading/trailing |)
            parts = [p for p in parts if p]
            
            if len(parts) >= 3:
                skill = parts[0]
                skill_type = parts[1]
                experience = parts[2]
                
                # Extract years from experience with better parsing
                years = ""
                if experience and experience != "N/A":
                    # Look for numbers in experience
                    years_match = re.search(r'(\d+\.?\d*)\s*(?:year|yr|years|y)\w*', experience, re.IGNORECASE)
                    if years_match:
                        years = years_match.group(1)
                    else:
                        # Try to find any number
                        any_num_match = re.search(r'(\d+)', experience)
                        if any_num_match:
                            years = any_num_match.group(1)
                
                # Map to proper type for VA template
                requirement_type = "Required"
                if 'desired' in skill_type.lower() or 'preferred' in skill_type.lower():
                    requirement_type = "Desired"
                elif 'highly' in skill_type.lower():
                    requirement_type = "Highly desired"
                
                skills_data.append({
                    'skill': skill,
                    'type': requirement_type,
                    'experience': experience,
                    'years': years
                })
                print(f"    ✅ Found skill: {skill} - {requirement_type} - {experience}")
                
            elif len(parts) == 1:
                # Just a skill name without type/experience
                skill = parts[0]
                skills_data.append({
                    'skill': skill,
                    'type': 'Required',
                    'experience': '',
                    'years': ''
                })
                print(f"    ✅ Found skill: {skill} (no type/experience)")
                
        except Exception as e:
            print(f"    ❌ Error processing skill row: {e}")
            continue
    
    print(f"  Total skills parsed: {len(skills_data)}")
    
    return skills_data

def parse_regex_extracted_skills(skills_text):
    """
    Parse skills from regex-extracted format like:
    SKILLS TABLE:
    Manage vendor relationships... Required 3 Years
    Experience tracking SLAs... Required 3 Years
    """
    skills_data = []
    
    if not skills_text:
        return skills_data
    
    print(f"  📝 Parsing regex-extracted skills format")
    
    # Remove "SKILLS TABLE:" header if present
    if skills_text.startswith("SKILLS TABLE:"):
        skills_text = skills_text.replace("SKILLS TABLE:", "").strip()
    
    lines = skills_text.strip().split('\n')
    
    for line_num, line in enumerate(lines, 1):
        line = line.strip()
        if not line:
            continue
        
        # Skip empty lines or section headers
        if line.lower() in ['skills:', 'skills table:']:
            continue
        
        # Initialize defaults
        skill_desc = line
        req_type = "Required"
        years = ""
        
        try:
            # Try to extract the pattern: "... Required 3 Years"
            # Look for "Years" at the end of the line
            if " Years" in line:
                # Find the last occurrence of "Years"
                years_idx = line.rfind(" Years")
                if years_idx != -1:
                    # Get the part before "Years"
                    before_years = line[:years_idx].strip()
                    # Get the part after "Years" (should be empty or continuation)
                    after_years = line[years_idx + 6:].strip()
                    
                    # Extract the number before "Years"
                    # Go backwards from years_idx to find the number
                    # Look for pattern: "Required 3" or "Desired 5" etc.
                    type_year_match = re.search(r'(\bRequired\b|\bDesired\b|\bHighly desired\b|\bHIGHLY DESIRED\b)\s+(\d+)', before_years, re.IGNORECASE)
                    
                    if type_year_match:
                        req_type = type_year_match.group(1).strip()
                        years = type_year_match.group(2).strip()
                        
                        # Get skill description (everything before the requirement type)
                        req_type_start = before_years.rfind(req_type)
                        if req_type_start != -1:
                            skill_desc = before_years[:req_type_start].strip()
                        else:
                            skill_desc = before_years
                    
                    # Standardize requirement type
                    req_type_upper = req_type.upper()
                    if 'HIGHLY' in req_type_upper:
                        req_type = "HIGHLY DESIRED"
                    elif 'REQUIRED' in req_type_upper:
                        req_type = "Required"
                    elif 'DESIRED' in req_type_upper:
                        req_type = "Desired"
                    else:
                        req_type = "Required"
                        
                    # Clean up skill description
                    # Remove trailing dots
                    while skill_desc.endswith('..') or skill_desc.endswith('.'):
                        skill_desc = skill_desc.rstrip('.')
                    skill_desc = skill_desc.strip()
                    
            # Alternative pattern matching for cases where above doesn't work
            else:
                # Try regex pattern: "skill.. Required 3 Years"
                pattern = r'(.+?)(?:\.\.|\.)\s+(Required|Desired|Highly desired|HIGHLY DESIRED)\s+(\d+)\s+Years'
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    skill_desc = match.group(1).strip()
                    req_type = match.group(2).strip()
                    years = match.group(3).strip()
                    
                    # Standardize
                    if req_type.upper() == 'HIGHLY desired'.upper():
                        req_type = "HIGHLY DESIRED"
                    elif 'required' in req_type.lower():
                        req_type = "Required"
                    elif 'desired' in req_type.lower():
                        req_type = "Desired"
        
        except Exception as e:
            print(f"    ⚠️ Error parsing line {line_num}: {e}")
            # Keep defaults
        
        # Final cleanup
        skill_desc = skill_desc.rstrip('.').strip()
        
        skills_data.append({
            'skill': skill_desc,
            'type': req_type,
            'experience': f"{years} Years",
            'years': years
        })
    
    print(f"  ✅ Parsed {len(skills_data)} skills from regex format")
    
    # Debug: Show parsed skills
    if skills_data:
        print(f"  📋 First 3 parsed skills:")
        for i, skill in enumerate(skills_data[:3]):
            print(f"    {i+1}. '{skill['skill']}' - {skill['type']} - {skill['years']} years")
    
    return skills_data

def extract_deadline_date(text):
    """
    Extract deadline date - ONLY look for "No New Submittals After"
    If not present, calculate 4 BUSINESS DAYS from today
    Returns MMDD format string
    """
    from datetime import datetime, timedelta
    import re
    
    print(f"  🔍 Looking for 'No New Submittals After' date...")
    
    # ONLY look for "No New Submittals After" patterns
    patterns = [
        # Pattern 1: "No New Submittals After: 09-16-2024"
        r'No New Submittals After:\s*(\d{1,2})-(\d{1,2})-(\d{4})',
        # Pattern 2: "No New Submittals After: 09/16/2024"
        r'No New Submittals After:\s*(\d{1,2})/(\d{1,2})/(\d{4})',
        # Pattern 3: "No New Submittals After: September 16, 2024"
        r'No New Submittals After:\s*([A-Za-z]+)\s+(\d{1,2}),\s+(\d{4})',
        # Pattern 4: "No New Submittals After: 09/16" (NEW - without year)
        r'No New Submittals After:\s*(\d{1,2})/(\d{1,2})\b',
        # Pattern 5: "No New Submittals After: 09-16" (NEW - without year)
        r'No New Submittals After:\s*(\d{1,2})-(\d{1,2})\b',
    ]
    
    # ==============================================
    # LOGIC: Check if phrase is PRESENT or NOT
    # ==============================================
    
    # Check if "No New Submittals After" exists anywhere in text
    if 'No New Submittals After' not in text:
        print(f"  ⚠️  'No New Submittals After' NOT PRESENT in text")
        print(f"  📅 Going to 4 BUSINESS DAYS calculation (Virginia logic)")
        return calculate_virginia_deadline()
    
    # If we get here, "No New Submittals After" IS PRESENT in text
    print(f"  ✅ 'No New Submittals After' IS PRESENT in text")
    
    # Try to extract the date using patterns
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            print(f"  ✅ Successfully extracted date from 'No New Submittals After'")
            try:
                # For numeric dates with year (09-16-2024 or 09/16/2024)
                if pattern in [patterns[0], patterns[1]]:
                    month = match.group(1).zfill(2)  # "9" -> "09"
                    day = match.group(2).zfill(2)    # "16" -> "16"
                    year = match.group(3)
                    print(f"  📅 Extracted deadline: {month}/{day}/{year}")
                    return f"{month}{day}"  # Return "0916"
                
                # For text month dates (September 16, 2024)
                elif pattern == patterns[2]:
                    month_name = match.group(1)
                    day = match.group(2).zfill(2)
                    year = match.group(3)
                    
                    # Convert month name to number
                    month_dict = {
                        'january': '01', 'february': '02', 'march': '03', 'april': '04',
                        'may': '05', 'june': '06', 'july': '07', 'august': '08',
                        'september': '09', 'october': '10', 'november': '11', 'december': '12',
                        'jan': '01', 'feb': '02', 'mar': '03', 'apr': '04',
                        'may': '05', 'jun': '06', 'jul': '07', 'aug': '08',
                        'sep': '09', 'oct': '10', 'nov': '11', 'dec': '12'
                    }
                    
                    month_lower = month_name.lower()
                    if month_lower in month_dict:
                        month = month_dict[month_lower]
                        print(f"  📅 Extracted deadline: {month}/{day}/{year}")
                        return f"{month}{day}"
                        
                # For dates WITHOUT year (12/11 or 12-11) - NEW PATTERNS
                elif pattern in [patterns[3], patterns[4]]:
                    month = match.group(1).zfill(2)  # "12" -> "12"
                    day = match.group(2).zfill(2)    # "11" -> "11"
                    # Get current year for display
                    current_year = datetime.now().year
                    print(f"  📅 Extracted deadline without year: {month}/{day} (assuming {current_year})")
                    return f"{month}{day}"  # Return "1211"
                        
            except Exception as e:
                print(f"  ⚠️  Error extracting date: {e}")
                continue
    
    # If we get here, "No New Submittals After" is PRESENT but date couldn't be extracted
    print(f"  ⚠️  'No New Submittals After' found but COULD NOT EXTRACT DATE")
    print(f"  📅 Going to 4 BUSINESS DAYS calculation")
    return calculate_virginia_deadline()

def calculate_virginia_deadline():
    """
    Calculate deadline for Virginia locations ONLY - 4 business days from today
    Should only be called when NO date is found in the content
    """
    from datetime import datetime, timedelta

    today = datetime.now()
    days_added = 0
    target_days = 4

    print(f"  📅 Calculating Virginia deadline (4 business days from today)")
    print(f"  📅 Today: {today.strftime('%A, %Y-%m-%d')}")

    # Start from tomorrow
    current_date = today + timedelta(days=1)
    
    while days_added < target_days:
        # Monday=0, Friday=4, Saturday=5, Sunday=6
        if current_date.weekday() < 5:  # Weekday (Mon-Fri)
            days_added += 1
        # Move to next day
        current_date += timedelta(days=1)
    
    # Go back one day (we overshot by 1)
    deadline_date = current_date - timedelta(days=1)
    deadline = deadline_date.strftime("%m%d")
    
    print(f"  📅 Calculated Virginia deadline: {deadline} ({deadline_date.strftime('%A, %B %d, %Y')})")
    return deadline

def extract_essential_data_with_regex(raw_content):
    """
    Extract all essential information from raw scraped data using regex
    Returns cleaned, structured dictionary with proper formatting
    """
    extracted = {}
    
    # === TITLE/ROLE EXTRACTION ===
    title_match = re.search(r'Title/Role:\s*(.+)', raw_content)
    if title_match:
        extracted['title'] = title_match.group(1).strip()
        print(f"  📝 Extracted Title/Role: {extracted['title']}")
    else:
        # Try alternative patterns if "Title/Role:" not found
        alt_title_patterns = [
            r'Title:\s*(.+)',
            r'Role:\s*(.+)', 
            r'Job Title:\s*(.+)',
            r'Position Title:\s*(.+)'
        ]
        
        for pattern in alt_title_patterns:
            match = re.search(pattern, raw_content)
            if match:
                extracted['title'] = match.group(1).strip()
                print(f"  📝 Extracted Title (alternative pattern): {extracted['title']}")
                break
        
        if 'title' not in extracted:
            extracted['title'] = "Position"
            print("  ⚠️  No title found, using default: 'Position'")
    
    # === WORKSITE ADDRESS ===
    # Capture multi-line address (until next section or blank line)
    worksite_match = re.search(r'Worksite Address:\s*(.*?)(?=\n\s*(?:CAI|Max Submittals|Per Opening|Currently Engaged|Expenses Allowed|Work Arrangement|Title/Role):|\n\n|\Z)', 
                              raw_content, re.DOTALL | re.IGNORECASE)
    
    if worksite_match:
        worksite = worksite_match.group(1).strip()
        
        # If there are newlines in the address, join them with space
        if '\n' in worksite:
            # Split by lines, strip each line, then join with space
            lines = [line.strip() for line in worksite.split('\n') if line.strip()]
            worksite = ' '.join(lines)
            print(f"  🏢 Extracted multi-line Worksite Address: {worksite}")
        else:
            print(f"  🏢 Extracted Worksite Address: {worksite}")
        
        extracted['worksite_address'] = worksite
    else:
        # Fallback: try single line pattern
        worksite_match = re.search(r'Worksite Address:\s*(.+?)(?=\n|$)', raw_content)
        if worksite_match:
            worksite = worksite_match.group(1).strip()
            extracted['worksite_address'] = worksite
            print(f"  🏢 Extracted Worksite Address (single line): {worksite}")
    
    # === JOB ID & BASIC INFO ===
    job_id_match = re.search(r'Job ID:\s*([A-Z]{2}-\d+)', raw_content)
    extracted['job_id'] = job_id_match.group(1) if job_id_match else "UNKNOWN"
    
    # === MAX SUBMITTALS BY VENDOR ===
    max_submittals_match = re.search(r'Max Submittals by Vendor:\s*(\d+)', raw_content)
    extracted['max_submittals'] = max_submittals_match.group(1) if max_submittals_match else "1"
    
    # === LOCATION ===
    # Simplified version that definitely works for your format
    work_location_match = re.search(r'Work\s*Location:\s*([A-Za-z0-9\-_]+)',raw_content)
    if work_location_match:
        work_location = work_location_match.group(1).strip()
        print(f"  🏢 Work Location: {work_location}")
    else:
        work_location = ""
        print("  ⚠️  Work Location not found")
    
    # === DATES & OPENINGS - EXTRACT ALL FIELDS ===
    # No. of Openings
    openings_match = re.search(r'No\. of Openings:\s*(\d+)', raw_content)
    extracted['openings'] = openings_match.group(1) if openings_match else "1"
    
    # Total No. Filled
    filled_match = re.search(r'Total No\. Filled:\s*(\d+)', raw_content)
    extracted['filled'] = filled_match.group(1) if filled_match else "0"
    
    # Start Date
    start_match = re.search(r'Start Date:\s*(\d{1,2}/\d{1,2}/\d{4})', raw_content)
    extracted['start_date'] = start_match.group(1) if start_match else ""
    
    # End Date
    end_match = re.search(r'End Date:\s*(\d{1,2}/\d{1,2}/\d{4})', raw_content)
    extracted['end_date'] = end_match.group(1) if end_match else ""
    
    # === DEADLINE EXTRACTION WITH VIRGINIA LOGIC ===
    deadline_result = extract_deadline_date(raw_content)
    extracted['deadline'] = deadline_result  # This will be in MMDD format
    
    # === WORK ARRANGEMENT ===
    work_arrangement_match = re.search(r'Work Arrangement:\s*([^\n]+)', raw_content)
    extracted['work_arrangement'] = work_arrangement_match.group(1).strip() if work_arrangement_match else "Onsite"
    
    # === BILL RATE EXTRACTION - FROM BOTH SECTIONS SEPARATELY ===
    bill_rates = extract_bill_rates_from_all_sections(raw_content)
    extracted.update(bill_rates)
    
    # === DESCRIPTIONS - PRESERVE ORIGINAL FORMATTING ===
    # Short Description - keep original formatting
    short_desc_match = re.search(r'SHORT DESCRIPTION:\s*(.*?)(?=COMPLETE DESCRIPTION:|\n\n|\Z)', raw_content, re.DOTALL)
    if short_desc_match:
        short_desc = short_desc_match.group(1).strip()
        # Remove extra whitespace but preserve line breaks
        short_desc = re.sub(r'[ \t]+', ' ', short_desc)  # Replace multiple spaces/tabs with single space
        short_desc = re.sub(r'\n[ \t]+\n', '\n\n', short_desc)  # Preserve paragraph breaks
        extracted['short_description'] = short_desc
    
    # Complete Description - keep original formatting  
    complete_desc_match = re.search(r'COMPLETE DESCRIPTION:\s*(.*?)(?=\n\n|\Z|===)', raw_content, re.DOTALL)
    if complete_desc_match:
        complete_desc = complete_desc_match.group(1).strip()
        # Remove extra whitespace but preserve line breaks
        complete_desc = re.sub(r'[ \t]+', ' ', complete_desc)  # Replace multiple spaces/tabs with single space
        complete_desc = re.sub(r'\n[ \t]+\n', '\n\n', complete_desc)  # Preserve paragraph breaks
        # Remove manager notes if present
        complete_desc = re.sub(r'MANAGER NOTES:.*?Description:', '', complete_desc, flags=re.DOTALL)
        extracted['complete_description'] = complete_desc
    
    # === SKILLS TABLE - CONVERT TO BULLET POINT FORMAT ===
    skills_section_match = re.search(r'=== SKILLS TABLE ===(.*?)(?=\n\n|\Z|===)', raw_content, re.DOTALL)
    if skills_section_match:
        skills_table = skills_section_match.group(1).strip()
    
        bullet_points = []
        for line in skills_table.split('\n'):
            line = line.strip()
            # Process only table data rows (skip headers and separators)
            if line.startswith('|') and '---' not in line and 'Skill' not in line:
                cells = [cell.strip() for cell in line.split('|')[1:-1] if cell.strip()]
                if len(cells) >= 3:
                    bullet_points.append(f"{cells[0]}. {cells[1]} {cells[2]}")
    
        extracted['skills_table'] = "\n".join(bullet_points) if bullet_points else ""
    
    return extracted

def extract_bill_rates_from_all_sections(raw_content):
    """
    Extract bill rates from ALL sections
    """
    print("=" * 60)
    print("🔍 DEBUG BILL RATE EXTRACTION")
    print("=" * 60)
    
    # Print a sample to see what's in the content
    print("📋 First 1000 chars of content:")
    print(raw_content[:1000])
    print("=" * 60)
    
    bill_rates = {
        'current_budget_rate': "00",
        'questions_rate': "00", 
        'short_desc_rate': "00",
        'final_bill_rate': "00",
        'source': 'none'
    }
    
    # === 1. Check for ANY dollar amounts in ENTIRE content ===
    print("💰 Looking for ANY $ amounts in entire content...")
    all_dollar_matches = re.findall(r'\$(\d+\.?\d+)', raw_content)
    print(f"   Found: {all_dollar_matches}")
    
    # === 2. Specifically check Questions section ===
    print("\n🔍 Looking for Questions section...")
    
    # Try different patterns to find Questions
    questions_patterns = [
        r'=== QUESTIONS ===(.*?)(?===|\Z)',
        r'QUESTIONS(.*?)(?=\n\n|\Z)',
        r'Q1:(.*?)(?=Q\d+:|$)',  # This pattern was causing issues, let's fix it
    ]
    
    questions_content = ""
    for pattern in questions_patterns:
        match = re.search(pattern, raw_content, re.DOTALL | re.IGNORECASE)
        if match:
            questions_content = match.group(1)
            print(f"✅ Found Questions section with pattern: {pattern[:30]}...")
            print(f"   Questions content length: {len(questions_content)} chars")
            print(f"   First 200 chars: {questions_content[:200]}")
            break
    
    if questions_content:
        print("\n💰 Searching for $ in Questions...")
        question_dollar_matches = re.findall(r'\$(\d+\.?\d+)', questions_content)
        print(f"   Found in Questions: {question_dollar_matches}")
        
        # Filter reasonable rates
        valid_rates = []
        for rate in question_dollar_matches:
            try:
                rate_num = float(rate)
                if 10 <= rate_num <= 200:
                    valid_rates.append(rate_num)
                    print(f"   ✅ Valid rate: ${rate_num}")
            except:
                continue
        
        if valid_rates:
            highest_rate = max(valid_rates)
            bill_rates['questions_rate'] = f"{highest_rate:.2f}"
            print(f"🎯 SET Questions rate to: ${bill_rates['questions_rate']}")
    
    # === 3. Check Current Budget ===
    print("\n💰 Looking for Current Budget $ amounts...")
    # Look for patterns like: $73.79 USD
    usd_matches = re.findall(r'\$(\d+\.?\d+)\s*USD', raw_content)
    print(f"   USD matches: {usd_matches}")
    
    valid_rates = []
    for rate in usd_matches:
        try:
            rate_num = float(rate)
            if 10 <= rate_num <= 200:
                valid_rates.append(rate_num)
                print(f"   ✅ Valid USD rate: ${rate_num}")
        except:
            continue
    
    if valid_rates:
        highest_rate = max(valid_rates)
        bill_rates['current_budget_rate'] = f"{highest_rate:.2f}"
        print(f"🎯 SET Current Budget rate to: ${bill_rates['current_budget_rate']}")
    
    # === Determine final rate ===
    print("\n" + "=" * 60)
    print("📊 RATE DECISION:")
    
    q_rate = float(bill_rates['questions_rate']) if bill_rates['questions_rate'] != "00" else 0
    c_rate = float(bill_rates['current_budget_rate']) if bill_rates['current_budget_rate'] != "00" else 0
    
    print(f"   Questions rate: ${q_rate}")
    print(f"   Current Budget rate: ${c_rate}")
    
    # Priority: Questions > Current Budget
    if q_rate > 0:
        bill_rates['final_bill_rate'] = bill_rates['questions_rate']
        bill_rates['source'] = 'questions'
        print(f"🏆 SELECTED: Questions rate ${bill_rates['questions_rate']}")
    elif c_rate > 0:
        bill_rates['final_bill_rate'] = bill_rates['current_budget_rate']
        bill_rates['source'] = 'current_budget'
        print(f"🥈 SELECTED: Current Budget rate ${bill_rates['current_budget_rate']}")
    else:
        bill_rates['final_bill_rate'] = "00"
        bill_rates['source'] = 'none'
        print("❌ NO RATE FOUND - using $00")
    
    print("=" * 60)
    return bill_rates

def save_extracted_data_to_file(file_path, extracted_data):
    """
    Save extracted data in the CLEAN FORMAT with separate bill rates
    """
    import os
    
    # Extract requisition number from filename
    requisition_number = "unknown"
    filename = os.path.basename(file_path)
    if "requisition_" in filename and "_complete.txt" in filename:
        try:
            parts = filename.split('_')
            if len(parts) >= 2 and parts[1].isdigit():
                requisition_number = parts[1]
        except:
            pass
    
    # Build the clean output format
    output_lines = []
    
    # Header information
    output_lines.append(f"Requisition Number: {requisition_number}")
    output_lines.append(f"Work Arrangement: {extracted_data.get('work_arrangement', 'Onsite')}")
    
    # ADD TITLE HERE
    output_lines.append(f"Title/Role: {extracted_data.get('title', 'Position')}")
    
    # Worksite Address - NEW (preserve formatting if it was multi-line)
    if extracted_data.get('worksite_address'):
        worksite = extracted_data['worksite_address']
        # If the address contains " | " separator, keep it as is
        if ' | ' in worksite:
            output_lines.append(f"Worksite Address: {worksite}")
        else:
            # For multi-line addresses joined with space, keep as single line
            output_lines.append(f"Worksite Address: {worksite}")
    
    # Location
    location_parts = []
    if extracted_data.get('worksite'):
        location_parts.append(extracted_data['worksite'])
    if extracted_data.get('work_location') and extracted_data['work_location'] not in ['GL:', 'N/A']:
        location_parts.append(extracted_data['work_location'])
    
    if location_parts:
        output_lines.append(f"Location: {' | '.join(location_parts)}")
    
    # All the fields
    output_lines.append(f"No. of Openings: {extracted_data.get('openings', '1')}")
    output_lines.append(f"Max Submittals by Vendor: {extracted_data.get('max_submittals', '1')}")
    output_lines.append(f"Total No. Filled: {extracted_data.get('filled', '0')}")
    output_lines.append(f"Start Date: {extracted_data.get('start_date', '')}")
    output_lines.append(f"End Date: {extracted_data.get('end_date', '')}")
    
    # Deadline - format as MM/DD if in MMDD format
    deadline = extracted_data.get('deadline', '')
    if deadline and len(deadline) == 4 and deadline.isdigit():
        # Format from MMDD to MM/DD
        formatted_deadline = f"{deadline[:2]}/{deadline[2:]}"
    else:
        formatted_deadline = deadline
    output_lines.append(f"No New Submittals After: {formatted_deadline}")
    
    output_lines.append(f"Current Budget Rate: ${extracted_data.get('current_budget_rate', '00')}/hour")
    output_lines.append(f"Questions Section Rate: ${extracted_data.get('questions_rate', '00')}/hour")
    output_lines.append(f"Short Description Rate: ${extracted_data.get('short_desc_rate', '00')}/hour")
    output_lines.append(f"Final Bill Rate: ${extracted_data.get('final_bill_rate', '00')}/hour (from {extracted_data.get('source', 'unknown')})")
    
    # Descriptions
    if extracted_data.get('short_description'):
        output_lines.append("SHORT DESCRIPTION:")
        output_lines.append(extracted_data['short_description'])
        output_lines.append("")  # Empty line
    
    if extracted_data.get('complete_description'):
        output_lines.append("COMPLETE DESCRIPTION:")
        output_lines.append(extracted_data['complete_description'])
        output_lines.append("")  # Empty line
    
    # Skills Table
    if extracted_data.get('skills_table'):
        output_lines.append("SKILLS TABLE:")
        output_lines.append(extracted_data['skills_table'])
    
    # Save to file
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(output_lines))
    
    print(f"  💾 Saved clean formatted data to: {os.path.basename(file_path)}")
    return '\n'.join(output_lines)

def process_single_file_regex_extraction(file_path):
    """
    Process a single file: RAW → REGEX EXTRACT → CLEAN FORMAT
    """
    try:
        # Read original scraped data
        with open(file_path, 'r', encoding='utf-8') as f:
            original_content = f.read()
        
        original_size = len(original_content)
        filename = os.path.basename(file_path)
        
        print(f"  🔍 Extracting: {filename}")
        print(f"    📊 Original: {original_size:,} chars")
        
        # Skip if already extracted (check for new format)
        if "Requisition Number:" in original_content:
            print("    ⏩ Already extracted with clean format, skipping")
            return True
        
        # Extract data with regex
        extracted_data = extract_essential_data_with_regex(original_content)
        
        # Save in CLEAN FORMAT (not JSON)
        clean_content = save_extracted_data_to_file(file_path, extracted_data)
        
        cleaned_size = len(clean_content)
        reduction = ((original_size - cleaned_size) / original_size) * 100
        
        print(f"    ✅ Reduced by: {reduction:.1f}%")
        print(f"    📁 File updated with clean format: {filename}")
        
        return True
        
    except Exception as e:
        print(f"    ❌ Error extracting {os.path.basename(file_path)}: {str(e)}")
        return False

def extract_all_scraped_files(output_dir='vms_outputs'):
    """
    Extract ALL scraped files: RAW → REGEX EXTRACT → SAME FILE
    This runs immediately after scraping is complete
    """
    output_dir = "vms_outputs"
    
    # Get all requisition files
    requisition_files = [os.path.join(output_dir, f) for f in os.listdir(output_dir) 
                        if f.endswith('.txt') and f.startswith('requisition_')]
    
    if not requisition_files:
        print("No scraped files found for extraction")
        return 0
    
    print(f"\n🔍 REGEX EXTRACTION: {len(requisition_files)} files")
    print("=" * 50)
    
    processed_count = 0
    
    for i, file_path in enumerate(requisition_files):
        success = process_single_file_regex_extraction(file_path)
        if success:
            processed_count += 1
        
        # Add small delay between files
        if i < len(requisition_files) - 1:
            time.sleep(1)
    
    print("=" * 50)
    print(f"✅ EXTRACTION COMPLETE: {processed_count}/{len(requisition_files)} files")
    
    return processed_count

def extract_all_skills_from_requisition(content):
    """
    Extract ALL skills from requisition content text
    Returns list of skill strings
    """
    skills = []
    
    if not content:
        return skills
    
    # Try multiple patterns to find the skills section
    patterns = [
        r'^skills:\s*\n(.*?)(?=\n\n|\n===|\nDescription:|\nJob ID:|\nLocation:|\nDuration:|\nPositions:|$)',
        r'^SKILLS:\s*\n(.*?)(?=\n\n|\n===|\nDescription:|\nJob ID:|\nLocation:|\nDuration:|\nPositions:|$)',
        r'skills:\s*\n(.*?)(?=\n\n|\n===|\nDescription:|\nJob ID:|\nLocation:|\nDuration:|\nPositions:|$)',
    ]
    
    skills_text = ""
    for pattern in patterns:
        match = re.search(pattern, content, re.DOTALL | re.MULTILINE | re.IGNORECASE)
        if match:
            skills_text = match.group(1).strip()
            break
    
    if not skills_text:
        # If no skills section found, return empty
        return skills
    
    # Parse skills from the text
    lines = skills_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for bullet points (•, -, *, etc.)
        if line.startswith(('•', '-', '*', '○', '∙', '·')):
            # Remove bullet point
            clean_line = re.sub(r'^[•\-*○∙·\s]+', '', line)
            clean_line = clean_line.strip()
            if clean_line:
                skills.append(clean_line)
        # Check for numbered items (1., 2., etc.)
        elif re.match(r'^\d+[\.\)]\s+', line):
            # Remove number
            clean_line = re.sub(r'^\d+[\.\)]\s+', '', line)
            clean_line = clean_line.strip()
            if clean_line:
                skills.append(clean_line)
        # Check if line looks like a skill (contains experience/years keywords)
        elif any(keyword in line.lower() for keyword in ['experience', 'knowledge', 'familiarity', 'ability', 'skill', 'proficiency', 'certification', 'years']):
            skills.append(line)
    
    return skills

def extract_title_from_requisition(content):
    """
    Extract the job title from requisition content and clean it up
    """
    # First, try to find the specific "Title/Role:" pattern
    title_match = re.search(r'Title/Role:\s*(.+)$', content, re.MULTILINE | re.IGNORECASE)
    if title_match:
        title = title_match.group(1).strip()
        # Clean up the title - remove any prefix like "NC FAST Requisition Class: DEV : "
        clean_title = re.sub(r'^.*?(?:requisition class|nc fast)[^:]*:\s*', '', title, flags=re.IGNORECASE)
        clean_title = clean_title.strip()
        return clean_title
    
    # Fallback: Look for patterns that might indicate a title
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        # Skip empty lines and section headers
        if not line or line.startswith('==') or len(line) > 100:
            continue
            
        # Common title indicators (case insensitive)
        if any(keyword in line.lower() for keyword in 
               ['engineer', 'developer', 'analyst', 'manager', 'specialist', 
                'consultant', 'architect', 'administrator', 'designer']):
            # Clean up the title - remove any prefix like "NC FAST Requisition Class: DEV : "
            clean_title = re.sub(r'^.*?(?:requisition class|nc fast)[^:]*:\s*', '', line, flags=re.IGNORECASE)
            clean_title = clean_title.strip()
            return clean_title
                
    # Final fallback: return the first meaningful line
    for line in lines:
        if line.strip() and not line.startswith('=='):
            clean_title = re.sub(r'^.*?(?:requisition class|nc fast)[^:]*:\s*', '', line.strip(), flags=re.IGNORECASE)
            return clean_title.strip()
    
    return "Position"

def update_rtr_document(req_id, title, state_abbr, documents_dir, output_path, requisition_content=None, skills_data=None):
    """
    Update the RTR document for the specific state with the requisition title
    ENHANCED: Handles combined documents like RTR_SM_Idaho.docx
    """
    # If state_abbr is DEFAULT or unclear, try to extract from requisition content
    if state_abbr == 'DEFAULT' and requisition_content:
        state_abbr = extract_state_from_job_id(requisition_content)
        print(f"  Extracted state from content: {state_abbr}")

    # Get state info - this defines what template we WANT to find
    state_info = STATE_MAPPING.get(state_abbr, STATE_MAPPING['DEFAULT'])
    desired_template_name = state_info['rtr_template']
    print(f"  Target state: {state_abbr}, Searching for template: {desired_template_name}")

    # Build the full desired path
    template_path = os.path.join(documents_dir, desired_template_name)
    
    # Check if the desired state template exists
    if os.path.exists(template_path):
        print(f"✅ Found exact state template: {desired_template_name}")
    else:
        print(f"❌ Desired template not found: {desired_template_name}")
        # Look for ANY template that matches the state abbreviation pattern
        all_files = os.listdir(documents_dir)
        
        # Filter out files that contain numbers (these are generated output files)
        template_files = [f for f in all_files if f.endswith('.docx') and not any(char.isdigit() for char in f)]
        print(f"  Available template files (no numbers): {[f for f in template_files if 'RTR' in f or 'SM' in f]}")
        
        # Look for template files that contain the state abbreviation
        matching_templates = []
        for file in template_files:
            file_upper = file.upper()
            # Match patterns like: RTR_FL.docx, SM_FL.docx, RTR_SM_Idaho.docx (NO NUMBERS)
            if (state_abbr.upper() in file_upper and
                not any(char.isdigit() for char in file)):
                matching_templates.append(file)
        
        if matching_templates:
            # Found a template for the correct state! Use the first match
            template_path = os.path.join(documents_dir, matching_templates[0])
            print(f"✅ Using found state-specific template: {matching_templates[0]}")
        else:
            # LAST RESORT: Use the default template
            default_template = STATE_MAPPING['DEFAULT']['rtr_template']
            template_path = os.path.join(documents_dir, default_template)
            if os.path.exists(template_path):
                print(f"⚠️  No state template found. Using default: {default_template}")
            else:
                print(f"❌ ERROR: No default template found either!")
                return None

    # Load the document
    try:
        doc = Document(template_path)
        print(f"✅ Loaded template: {os.path.basename(template_path)}")
    except Exception as e:
        print(f"❌ Error loading template {template_path}: {e}")
        return None

    # === FIX: UNIVERSAL SKILLS DATA HANDLING FOR ALL STATES ===
    # If skills_data is None but we have requisition_content, try to extract skills
    if skills_data is None and requisition_content:
        print("  ⚠️ No skills data provided, attempting to extract from content...")
        if "=== SKILLS TABLE ===" in requisition_content:
            try:
                skills_part = requisition_content.split("=== SKILLS TABLE ===")[1]
                skills_table_content = skills_part.split("\n\n")[0].strip()
                skills_data = parse_skills_table(skills_table_content)
                print(f"  ✅ Extracted {len(skills_data)} skills from requisition content")
            except Exception as e:
                print(f"  ❌ Failed to extract skills from content: {e}")
                skills_data = []
        else:
            print("  ❌ No skills table found in requisition content")
            skills_data = []
    elif skills_data is None:
        print("  ❌ No skills data available and no content to extract from")
        skills_data = []

    # Format the title consistently for both places
    formatted_title = f"{title} ({req_id})"
    state_name = state_info['full_name']
    print(f"  Updating RTR document for {state_name} with title: {formatted_title}")

    # Update all state references in the document - CAREFUL REPLACEMENT
    state_updated = False
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        
        # Only replace state references that are clearly template text, not job data
        is_template_text = any(pattern in original_text for pattern in [
            "Managed Services Provider Contract",
            "VectorVMS Requirement Number and Title",
            "Candidate Full Legal Name",
            "Candidate Pay Rate for this Position",
            "Candidate Employment Type if Selected for Engagement",
            "has the sole right to represent me in matters of work assignment"
        ])
        
        if is_template_text:
            # Replace state-specific content - Only replace if it's NOT the correct state
            for target_state_abbr, target_state_info in STATE_MAPPING.items():
                if target_state_abbr == 'DEFAULT' or target_state_abbr == state_abbr:
                    continue  # Skip DEFAULT and current state
                    
                # Replace full state names from other states (only in template text)
                if target_state_info['full_name'] in original_text:
                    new_text = original_text.replace(target_state_info['full_name'], state_name)
                    paragraph.text = new_text
                    state_updated = True
                    print(f"   Replaced '{target_state_info['full_name']}' with '{state_name}' in template text")
                
                # Replace state abbreviations from other states (only in template text)
                elif f" {target_state_abbr} " in f" {original_text} ":
                    new_text = original_text.replace(f" {target_state_abbr} ", f" {state_abbr} ")
                    paragraph.text = new_text
                    state_updated = True
                    print(f"   Replaced '{target_state_abbr}' with '{state_abbr}' in template text")
    
    if not state_updated:
        print("   ⚠️ No state-specific template content found to replace")
    
    # === FIX: TITLE CONSISTENCY ===
    # Update the title in email subject
    subject_updated = False
    email_subject_found = False
    for paragraph in doc.paragraphs:
        if "INSERT THE FOLLOWING INTO EMAIL SUBJECT" in paragraph.text:
            email_subject_found = True
        elif email_subject_found and paragraph.text.strip():
            # Found the subject line, update it with consistent title
            paragraph.text = formatted_title
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(14)
                paragraph.runs[0].bold = True
            print(f"   Updated email subject: {formatted_title}")
            subject_updated = True
            break
    
    # Update the title in Vector VMS section with the SAME title
    vms_updated = False
    vms_section_found = False
    for paragraph in doc.paragraphs:
        # KEEPING BOTH TEXT PATTERNS: "VectorVMS" AND "Vector VMS"
        if "VectorVMS Requirement Number and Title" in paragraph.text or "Vector VMS Requirement Number and Title" in paragraph.text:
            vms_section_found = True
        elif vms_section_found and paragraph.text.strip():
            # Found the VMS title, update it with the SAME consistent title
            paragraph.text = formatted_title
            if paragraph.runs:
                paragraph.runs[0].font.size = Pt(12)
            print(f"   Updated VMS title: {formatted_title}")
            vms_updated = True
            break

    # === FIX: UNIVERSAL SKILLS TABLE UPDATE FOR ALL STATES ===
    skills_updated_in_rtr = False
    
    # Check if we have skills data for ANY state (not just Florida)
    if skills_data and doc.tables:
        print(f"  Processing skills table for {state_abbr} document...")
        
        for table_idx, table in enumerate(doc.tables):
            # Look for ANY skills table with specific columns
            if len(table.rows) > 0:
                header_cells = table.rows[0].cells
                header_text = ' '.join(cell.text.strip().lower() for cell in header_cells)
                
                # UNIVERSAL table detection - look for skills-related headers
                is_skills_table = (
                    any('skill' in cell.text.lower() for cell in header_cells) or
                    any('required' in cell.text.lower() or 'desired' in cell.text.lower() for cell in header_cells) or
                    any('experience' in cell.text.lower() for cell in header_cells) or
                    len(header_cells) >= 3  # Most skills tables have at least 3 columns
                )
                
                if is_skills_table:
                    print(f"  Found skills table at index {table_idx} with {len(header_cells)} columns")
                    
                    # Remove existing data rows (keep header row only)
                    while len(table.rows) > 1:
                        table._tbl.remove(table.rows[1]._tr)
                    
                    # Add new skills from skills_data - UNIVERSAL FORMAT
                    skills_added_count = 0
                    for skill in skills_data:
                        if not skill or not skill.get('skill') or skill['skill'] == 'N/A':
                            continue
                        
                        # Add a new row for each skill
                        row_cells = table.add_row().cells
                        
                        # SAFE column access - handle different column counts
                        num_columns = len(row_cells)
                        
                        # Map data to appropriate columns based on table structure
                        if num_columns >= 1:
                            row_cells[0].text = skill.get('skill', '') or ""
                        
                        if num_columns >= 2:
                            # Map 'Required'/'Desired' based on skill type
                            skill_type = skill.get('type', '')
                            if 'required' in skill_type.lower():
                                row_cells[1].text = 'Required'
                            elif 'desired' in skill_type.lower() or 'highly' in skill_type.lower():
                                row_cells[1].text = 'Highly desired'
                            else:
                                row_cells[1].text = skill_type or ""
                        
                        if num_columns >= 3:
                            experience = skill.get('experience', '')
                            # Extract just the numeric part (e.g., "5 years" → "5")
                            years_match = re.search(r'(\d+)\s*(year|yr|years)?', experience, re.IGNORECASE)
                            years_value = years_match.group(1) if years_match else skill.get('years', '0')
                            row_cells[2].text = years_value
                        
                        # Handle additional columns if they exist
                        if num_columns >= 4:
                             # For "Last Used" column (can be left empty or use current year)
                            row_cells[4].text = "Current"  # or leave empty ""
                        
                        if num_columns >= 5:
                            # For "Last Used" column (can be left empty or use current year)
                            row_cells[4].text = "Current"  # or leave empty ""
                        
                        skills_added_count += 1
                    
                    print(f"    Added {skills_added_count} skills to the {state_abbr} RTR table.")
                    skills_updated_in_rtr = True
                    set_table_borders(table)
                    break

        if not skills_updated_in_rtr:
            print(f"  No skills table found within {state_abbr} RTR document.")
    else:
        print(f"  No skills data available for {state_abbr} document.")
    
    # ENHANCED: Handle combined document naming (like RTR_SM_Idaho)
    output_filename = ""
    template_basename = os.path.basename(template_path).upper()
    
    # Check if this is a combined document (contains both RTR and SM in name)
    if 'RTR' in template_basename and 'SM' in template_basename:
        output_filename = f"RTR_SM_{state_abbr}_{req_id}.docx"
        print(f"  📄 Detected combined document, using naming: {output_filename}")
    else:
        output_filename = f"RTR_{state_abbr}_{req_id}.docx"
    
    output_path = os.path.join(os.path.dirname(output_path), output_filename)
    
    # Save the document
    try:
        doc.save(output_path)
        print(f"✅ Saved RTR document: {os.path.basename(output_path)}")
        return output_path
    except Exception as e:
        print(f"❌ Error saving RTR document: {e}")
        return None

def update_sm_document(skills_data, state_abbr, documents_dir, output_path, requisition_content=None):
    """
    Update the SM document with skills data for the specific state
    ENHANCED: Handles combined documents like RTR_SM_Idaho.docx
    """
    # If state_abbr is DEFAULT or unclear, try to extract from requisition content
    if state_abbr == 'DEFAULT' and requisition_content:
        state_abbr = extract_state_from_job_id(requisition_content)
        print(f"  Extracted state from content: {state_abbr}")
    
    # Get the correct template based on state - support both .doc and .docx
    state_info = STATE_MAPPING.get(state_abbr, STATE_MAPPING['DEFAULT'])
    
    # ENHANCED: Check if this state uses a combined document
    sm_template_name = state_info['sm_template']
    rtr_template_name = state_info['rtr_template']
    
    # If both templates point to the same file, it's a combined document
    is_combined_document = (sm_template_name == rtr_template_name)
    
    if is_combined_document:
        print(f"  🔄 Detected combined document state: {state_abbr}")
        print(f"  Using combined template: {sm_template_name}")
        
        # For combined documents, we need to extract req_id from output_path
        req_id_match = re.search(r'(\d+)\.docx$', output_path)
        req_id = req_id_match.group(1) if req_id_match else "unknown"
        
        # The SM document is the same as the RTR document for combined states
        combined_filename = f"RTR_SM_{state_abbr}_{req_id}.docx"
        combined_path = os.path.join(os.path.dirname(output_path), combined_filename)
        
        # Check if the combined document was already created by update_rtr_document
        if os.path.exists(combined_path):
            print(f"  ✅ Using existing combined document for SM: {combined_filename}")
            return combined_path
        else:
            print(f"  ❌ Combined document not found for SM: {combined_filename}")
            return None
    
    # Existing separate document logic for non-combined states
    # First try .docx version
    template_path_docx = os.path.join(documents_dir, state_info['sm_template'])
    template_path_doc = os.path.join(documents_dir, state_info['sm_template'].replace('.docx', '.doc'))
    
    # Check which template exists
    if os.path.exists(template_path_docx):
        template_path = template_path_docx
    elif os.path.exists(template_path_doc):
        template_path = template_path_doc
        print(f"  Using .doc template instead of .docx")
    else:
        print(f"❌ STATE TEMPLATE NOT FOUND: {state_info['sm_template']} for {state_abbr}")
        
        # Try to find any SM template that might work for this state (both .doc and .docx)
        all_sm_templates = []
        for ext in ['.docx', '.doc']:
            all_sm_templates.extend([f for f in os.listdir(documents_dir) 
                                   if f.startswith('SM_') and f.endswith(ext)])
        
        # Prioritize templates that match the state abbreviation
        state_specific_templates = [f for f in all_sm_templates if state_abbr in f.upper()]
        
        if state_specific_templates:
            template_path = os.path.join(documents_dir, state_specific_templates[0])
            print(f"   Using state-specific template: {state_specific_templates[0]}")
        elif all_sm_templates:
            print(f"   Available SM templates: {all_sm_templates}")
            template_path = os.path.join(documents_dir, all_sm_templates[0])
            print(f"   Using alternative template: {all_sm_templates[0]}")
        else:
            # Fallback to default (try both extensions)
            default_docx = os.path.join(documents_dir, STATE_MAPPING['DEFAULT']['sm_template'])
            default_doc = os.path.join(documents_dir, STATE_MAPPING['DEFAULT']['sm_template'].replace('.docx', '.doc'))
            
            if os.path.exists(default_docx):
                template_path = default_docx
            elif os.path.exists(default_doc):
                template_path = default_doc
            else:
                print(f"❌ ERROR: No default template found either!")
                return None
    
    print(f"✅ Using SM template: {os.path.basename(template_path)} for state {state_abbr}")
    
    # Load the document
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"❌ Error loading template {template_path}: {e}")
        return None
    
    # Initialize num_columns with a default value to prevent UnboundLocalError
    num_columns = 0
    skills_added = 0
    
    # Update skills table - SAFELY handle different table structures
    if doc.tables:
        table = doc.tables[0]
        
        # Get number of columns from the first row
        if len(table.rows) > 0:
            num_columns = len(table.rows[0].cells)
        
        # Remove existing data rows (keep header row)
        for i in range(len(table.rows) - 1, 0, -1):
            table._tbl.remove(table.rows[i]._tr)
        
        # Add new skills - SAFELY handle different column counts
        for skill in skills_data:
            if not skill['skill'] or skill['skill'] == 'N/A':
                continue
                
            row_cells = table.add_row().cells
            
            # SAFE column access - check how many columns exist
            if num_columns >= 1:
                row_cells[0].text = skill['skill']
            if num_columns >= 2:
                row_cells[1].text = skill['type']
            if num_columns >= 3:
                row_cells[2].text = skill['years']
            if num_columns >= 4:
                row_cells[3].text = ""
            if num_columns >= 5:
                row_cells[4].text = ""
            
            skills_added += 1
        
        print(f"   Added {skills_added} skills to SM document (table has {num_columns} columns)")
        set_table_borders(table)
    else:
        print(f"   No tables found in SM document, no skills added")
    
    # Ensure correct output filename with state code (always save as .docx)
    output_path = output_path.replace("_ID_", f"_{state_abbr}_")
    if f"_{state_abbr}_" not in output_path:
        base_name = os.path.basename(output_path)
        if "SM_" in base_name:
            output_path = os.path.join(os.path.dirname(output_path), f"SM_{state_abbr}_{req_id}.docx")
        else:
            output_path = os.path.join(os.path.dirname(output_path), f"SM_{state_abbr}_{req_id}.docx")
    
    # Save the document
    try:
        doc.save(output_path)
        print(f"✅ Saved SM document: {os.path.basename(output_path)}")
        return output_path
    except Exception as e:
        print(f"❌ Error saving SM document: {e}")
        return None

def clear_updated_documents():
    """Clear the updated_documents directory before processing"""
    updated_docs_dir = "vms_documents"
    
    if not os.path.exists(updated_docs_dir):
        os.makedirs(updated_docs_dir)
        print(f"Created directory: {updated_docs_dir}")
    else:
        # Remove all files in the directory
        for file in os.listdir(updated_docs_dir):
            file_path = os.path.join(updated_docs_dir, file)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
                    print(f"Removed previous file: {file}")
            except Exception as e:
                print(f"Error deleting file {file_path}: {e}")
        print(f"Cleared output directory: {updated_docs_dir}")

def add_title_to_formatted_content(formatted_content, title):
    """Add the generated title to the formatted content"""
    if not title or not formatted_content:
        return formatted_content
    
    # Insert title at the beginning with proper spacing
    return f"{title}\n\n{formatted_content}"