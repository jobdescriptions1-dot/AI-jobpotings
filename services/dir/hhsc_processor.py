import os
import json
import base64
import re
import time
from datetime import datetime
import pythoncom
import pytz
import PyPDF2
import tempfile
import win32com.client
from pdf2docx import Converter
from docx.enum.text import WD_BREAK
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.keys import Keys
import requests
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from dotenv import load_dotenv
import threading
import logging

from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
import os

from .regex_extractor import extract_data_with_regex
from .llm_processor import process_files_with_llm
from .email_sender import HHSCOutlookEmailSender
from docx import Document


# Load environment variables
load_dotenv()


# Define DEPARTMENT_CREDENTIALS
DEPARTMENT_CREDENTIALS = {
    'username': os.getenv('DEPARTMENT_USERNAME'),  # Make sure this matches your .env variable name
    'password': os.getenv('DEPARTMENT_PASSWORD') 
}
class SolicitationAutomation:
    def __init__(self):
        self.is_running = False
        self.processed_emails = set()  # Track ALL processed emails
        self.check_interval = 300  # 5 minutes
        self.last_email_count = 0
    
    def clear_folders_for_new_batch(self):
        """Clear folders when starting to process NEW email batch"""
        print("\n🗑️  CLEARING FOLDERS FOR NEW EMAIL BATCH")
        
        # Clear dir_documents folder
        documents_dir = "dir_documents"
        if os.path.exists(documents_dir):
            files_cleared = 0
            for file in os.listdir(documents_dir):
                file_path = os.path.join(documents_dir, file)
                try:
                    if os.path.isfile(file_path):
                        os.remove(file_path)
                        files_cleared += 1
                        print(f"  ✅ Cleared: {file}")
                except Exception as e:
                    print(f"  ⚠️  Could not clear {file}: {e}")
            print(f"  📂 Cleared {files_cleared} files from {documents_dir}")
        else:
            os.makedirs(documents_dir)
            print(f"  📂 Created directory: {documents_dir}")
        
        # Clear dir_portal_outputs folder  
        output_dir = "dir_portal_outputs"
        if os.path.exists(output_dir):
            files_cleared = 0
            for file in os.listdir(output_dir):
                file_path = os.path.join(output_dir, file)
                try:
                    if os.path.isfile(file_path):
                        os.remove(file_path)
                        files_cleared += 1
                except:
                    pass
            print(f"  📂 Cleared {files_cleared} files from {output_dir}")
        else:
            os.makedirs(output_dir)
            print(f"  📂 Created directory: {output_dir}")
    
    def trigger_dual_table_for_dir(self):
        """Trigger dual table processing after DIR processing"""
        try:
            print("\n📊 TRIGGERING DUAL TABLE FOR DIR FILES")
            
            # Check if there are output files
            output_dir = "dir_portal_outputs"
            if os.path.exists(output_dir):
                txt_files = [f for f in os.listdir(output_dir) if f.endswith('.txt')]
                
                if txt_files:
                    print(f"  ✅ Found {len(txt_files)} DIR files, running dual table...")
                    
                    # Import and run dual table
                    from services.dual_table.dual_table_service import run_dual_table_processing
                    run_dual_table_processing()
                    
                    print("  ✅ Dual table completed for DIR files")
                else:
                    print("  ⏭️  No DIR output files found")
            else:
                print("  ⏭️  No DIR output directory")
                
        except Exception as e:
            print(f"  ❌ Error triggering dual table: {e}")

    def monitor_emails_intelligently(self):
        """Intelligent monitoring - only processes NEWLY RECEIVED emails"""
        print("🚀 Starting Intelligent Solicitation Automation Monitor")
        print("📧 Will process ONLY NEWLY RECEIVED emails (not previously processed)")
        print("🗑️  Folders will be cleared when new emails are found")
        print("⏰ Checking every 5 minutes")
        
        while self.is_running:
            try:
                current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                print(f"🔍 Checking for NEW emails at {current_time}")
                
                # Get ONLY NEW emails that haven't been processed yet
                new_emails = self.get_new_emails()
                
                if new_emails:
                    print(f"🎯 NEW EMAILS DETECTED! Processing {len(new_emails)} new email(s)")
                    
                    # ⭐⭐ CLEAR FOLDERS BEFORE PROCESSING NEW EMAILS
                    self.clear_output_folders()
                    
                    self.run_single_automation_cycle(new_emails)
                    print(f"✅ Processing completed. Continuing to monitor for NEW emails...")
                else:
                    print(f"📭 No NEW emails detected. Continuing to monitor...")
                
                # Wait before next check
                print(f"⏳ Next check in {self.check_interval} seconds...")
                time.sleep(self.check_interval)
                
            except Exception as e:
                print(f"💥 Error in monitoring: {e}")
                print("🔄 Retrying in 60 seconds...")
                time.sleep(60)
        
    
    def clear_output_folders(self):
        """Clear output folders before processing new emails"""
        print(f"🗑️  Clearing output folders for new emails...")
        output_dirs = ["dir_portal_outputs", "dir_documents"]
        
        for output_dir in output_dirs:
            if os.path.exists(output_dir):
                files_deleted = 0
                for file in os.listdir(output_dir):
                    file_path = os.path.join(output_dir, file)
                    try:
                        if os.path.isfile(file_path):
                            os.unlink(file_path)
                            files_deleted += 1
                    except Exception as e:
                        print(f"❌ Error deleting {file_path}: {e}")
                print(f"✅ Cleared {files_deleted} files from {output_dir}")
            else:
                os.makedirs(output_dir)
                print(f"✅ Created directory: {output_dir}")
        
    
    def get_new_emails(self):
        """Get ONLY NEW emails that haven't been processed yet"""
        try:
            gmail_service = auto_authenticate_primary_gmail()
            
            # Search for TODAY'S department emails
            today_date = datetime.now().strftime("%Y-%m-%d")
            messages = search_specific_hhsc_email(gmail_service, days_back=1)
            
            if not messages:
                return []
            
            # ⭐⭐ CRITICAL FIX: Filter ONLY UNPROCESSED emails
            new_messages = []
            for message in messages:
                # Check if this email ID has been processed
                if message['id'] not in self.processed_emails:
                    new_messages.append(message)
                    print(f"  ✅ NEW email found: {message.get('department', 'Unknown')}")
                else:
                    print(f"  ⏭️  SKIPPING - Already processed: {message.get('department', 'Unknown')}")
            
            print(f"📊 Email Status: {len(new_messages)} NEW, {len(messages) - len(new_messages)} already processed")
            return new_messages
            
        except Exception as e:
            print(f"❌ Error getting new emails: {e}")
            return []
        
    
    def run_single_automation_cycle(self, new_messages):
        """Run ONE complete automation cycle for ONLY NEW emails"""
        driver = None
        try:
            print(f"\n🔥 STARTING NEW AUTOMATION CYCLE FOR {len(new_messages)} NEW EMAILS")
            print(f"📅 Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            if not new_messages:
                print("❌ No new messages to process")
                return
            
            self.clear_folders_for_new_batch()


            # Initialize driver and Gmail service
            driver = initialize_driver()
            
            processed_portals = 0
            
            # Get the first portal link to login
            first_portal_link = None
            gmail_service = auto_authenticate_primary_gmail()
            for message in new_messages:
                email_details = get_email_full_content(gmail_service, message['id'])
                if email_details:
                    portal_link = extract_portal_link(email_details['full_body'])
                    if portal_link:
                        first_portal_link = portal_link
                        break
            
            if not first_portal_link:
                print("❌ No valid portal links found in new emails")
                if driver:
                    driver.quit()
                return
            
            # Login once
            print(f"\n=== SINGLE LOGIN FOR NEW DEPARTMENTS ===")
            login_success = login_to_hhsc_portal(driver, first_portal_link, DEPARTMENT_CREDENTIALS)
            
            if not login_success:
                print("❌ Failed to login to portal")
                if driver:
                    driver.quit()
                return
            
            print("✅ Successfully logged in! Processing NEW departments...")
            
            # Process all NEW departments with the same session
            for i, message in enumerate(new_messages, 1):
                print(f"\n📥 Processing NEW email {i}/{len(new_messages)}...")
                
                email_details = get_email_full_content(gmail_service, message['id'])
                department = message.get('department', 'Unknown Department')
                
                if not email_details:
                    print("❌ Failed to get email content")
                    continue
                    
                display_email_info(email_details, department)
                portal_link = extract_portal_link(email_details['full_body'])
                
                if portal_link:
                    print(f"🔗 Portal link: {portal_link}")
                    print(f"🏢 Department: {department}")
                    
                    # Process the portal
                    success = process_portal_without_login(driver, portal_link, department)
                    
                    if success:
                        processed_portals += 1
                        # ⭐⭐ MARK as processed to prevent future processing
                        self.processed_emails.add(message['id'])
                        print(f"✅ Successfully processed {department}")
                        
                        # Short delay between departments
                        if i < len(new_messages):
                            print("⏳ Quick preparation for next department...")
                            time.sleep(3)
                    else:
                        print(f"❌ Failed to process {department}")
                else:
                    print(f"❌ No portal link found in {department} email")
            
            # Logout once
            print(f"\n=== SINGLE LOGOUT ===")
            logout_success = logout_from_portal(driver)
            if logout_success:
                print("✅ Successfully logged out")
            else:
                print("⚠️  Logout had issues")
            
            print(f"\n🎯 NEW EMAILS PROCESSING COMPLETED!")
            print(f"   Successfully processed {processed_portals} out of {len(new_messages)} new solicitations")
            
            if processed_portals > 0:
                # Process documents
                process_all_downloaded_documents()
                
                # REGEX EXTRACTION - ADDED BEFORE LLM PROCESSING
                print(f"\n=== REGEX DATA EXTRACTION ===")
                regex_processed = extract_data_with_regex()
                print(f"✅ Regex extraction completed: {regex_processed} files processed")
                
                # LLM PROCESSING
                print(f"\n=== LLM DATA PROCESSING ===")
                process_files_with_llm()

                # EMAIL SENDING
                print(f"\n=== SENDING RESPONSE EMAILS ===")
                email_sender = HHSCOutlookEmailSender()
                emails_sent = email_sender.send_emails_for_all_solicitations()
                print(f"✅ Sent {emails_sent} response emails")
                
                self.trigger_dual_table_for_dir()
                
                print(f"\n🎉 CYCLE COMPLETED - CONTINUING TO MONITOR FOR NEW EMAILS")
                print(f"   Total processed today: {len(self.processed_emails)} emails")
                print(f"   🔄 Still monitoring for NEW incoming emails...")
            else:
                print("\n❌ No new portals were successfully processed.")
                print("🔄 Continuing to monitor for NEW emails...")
            
            if driver:
                driver.quit()
            
        except Exception as e:
            print(f"❌ Error in automation cycle: {e}")
            import traceback
            traceback.print_exc()
            if driver:
                try:
                    driver.quit()
                except:
                    pass
        pass
    
    def start_automation(self):
        """Start the intelligent automation"""
        self.is_running = True
        
        self.monitor_thread = threading.Thread(target=self.monitor_emails_intelligently)
        self.monitor_thread.daemon = True
        self.monitor_thread.start()
        print("✅ Intelligent Automation started!")
        print("📧 Will only process NEWLY RECEIVED emails")
        print("🗑️  Folders will be cleared for each new batch")
        print("🔄 Automation will continue running indefinitely...")
        pass
    
    def stop_automation(self):
        """Stop the automation"""
        self.is_running = False
        print("🛑 Stopping automation...")
        pass
    
    def get_status(self):
        """Get current automation status"""
        return {
            'is_running': self.is_running,
            'processed_count': len(self.processed_emails),
            'last_check': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'status': '🔄 Monitoring for NEW emails...'
        }
        pass

# ===== ORIGINAL FUNCTIONS FROM texas1.py =====

def auto_authenticate_primary_gmail():
    """Authenticates with primary Gmail account (token.json)"""
    print("\n=== Authenticating Primary Gmail Account ===")
    creds = None
    token_path = 'token.json'
    SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
   
    if os.path.exists(token_path):
        print("Found primary token file, loading credentials...")
        try:
            creds = Credentials.from_authorized_user_file(token_path, SCOPES)
            if creds.expired and creds.refresh_token:
                print("Primary credentials expired, refreshing...")
                creds.refresh(Request())
        except Exception as e:
            print(f"Error loading primary credentials: {e}")
            creds = None
   
    if not creds or not creds.valid:
        print("No valid primary credentials found, initiating OAuth flow...")
        try:
            flow = InstalledAppFlow.from_client_secrets_file('client.json', SCOPES)
            creds = flow.run_local_server(port=8080)
            token_data = json.loads(creds.to_json())
            token_data['creation_time'] = datetime.now(pytz.UTC).isoformat()
            with open(token_path, 'w') as token:
                json.dump(token_data, token)
            print("Primary authentication successful! Token saved.")
        except Exception as e:
            print(f"Primary authentication failed: {e}")
            raise  
    try:
        print("Building primary Gmail service...")
        gmail_service = build('gmail', 'v1', credentials=creds)
        print("Primary Gmail service ready!")
        return gmail_service
    except Exception as e:
        print(f"Failed to build primary Gmail service: {e}")
        raise  
    pass

def initialize_driver():
    """Initialize and return a Chrome WebDriver with optimal settings"""
    options = webdriver.ChromeOptions()
    options.add_argument("--start-maximized")
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_experimental_option("excludeSwitches", ["enable-automation"])
    options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36")
    
    # Create documents folder if it doesn't exist
    documents_dir = os.path.join(os.getcwd(), "dir_documents")
    if not os.path.exists(documents_dir):
        os.makedirs(documents_dir)
    
    # Set download preferences
    prefs = {
        "download.default_directory": documents_dir,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True,
        "profile.default_content_settings.popups": 0
    }
    options.add_experimental_option("prefs", prefs)
    
    # USE THIS SIMPLE APPROACH - no os_type parameter
    service = Service("chromedriver\\chromedriver.exe")
    driver = webdriver.Chrome(service=service, options=options)
    return driver
    pass

def login_to_hhsc_portal(driver, url, credentials, max_retries=3):
    """Logs into various department portals with enhanced reliability"""
    print(f"  Trying credentials for portal: {url}")
    print(f"  Username: {credentials['username']}")
    driver.get(url)
    time.sleep(5)
 
    for attempt in range(max_retries):
        try:
            WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.TAG_NAME, 'body')))
           
            username_field = None
            password_field = None
            
            # Expanded list of selectors for different portal layouts
            username_selectors = [
                "//input[@type='text' or @type='email' or @type='username']",
                "//input[contains(@id, 'username') or contains(@name, 'username') or contains(@placeholder, 'Username') or contains(@placeholder, 'Email') or contains(@placeholder, 'User')]",
                "//input[contains(@id, 'email') or contains(@name, 'email')]",
                "//input[@id='username']",
                "//input[@name='username']",
                "//input[@id='email']",
                "//input[@name='email']",
                "//input[@id='input-1']",  # Common Salesforce selector
                "//input[@id='input-2']",  # Common Salesforce selector
                "//input[contains(@class, 'username')]",
                "//input[contains(@class, 'email')]",
                "//input[@data-id='username']",
                "//input[@aria-label='Username' or @aria-label='Email']"
            ]
            
            password_selectors = [
                "//input[@type='password']",
                "//input[contains(@id, 'password') or contains(@name, 'password')]",
                "//input[@id='password']",
                "//input[@name='password']",
                "//input[@id='input-3']",  # Common Salesforce selector
                "//input[@id='input-4']",  # Common Salesforce selector
                "//input[contains(@class, 'password')]",
                "//input[@data-id='password']",
                "//input[@aria-label='Password']"
            ]
            
            # Try to find username field
            for selector in username_selectors:
                try:
                    username_field = WebDriverWait(driver, 3).until(
                        EC.presence_of_element_located((By.XPATH, selector)))
                    if username_field and username_field.is_displayed():
                        print(f"  Found username field with selector: {selector}")
                        break
                except:
                    continue
            
            # Try to find password field
            for selector in password_selectors:
                try:
                    password_field = WebDriverWait(driver, 3).until(
                        EC.presence_of_element_located((By.XPATH, selector)))
                    if password_field and password_field.is_displayed():
                        print(f"  Found password field with selector: {selector}")
                        break
                except:
                    continue
            
            # If standard selectors didn't work, try more aggressive search
            if not username_field or not password_field:
                print("  Could not find login fields with standard selectors, trying alternative methods...")
                
                # Take screenshot for debugging
                try:
                    screenshot_path = f"debug_login_attempt_{attempt}.png"
                    driver.save_screenshot(screenshot_path)
                    print(f"  📸 Saved screenshot: {screenshot_path}")
                except:
                    pass
                
                # Get all input fields and analyze them
                all_inputs = driver.find_elements(By.TAG_NAME, "input")
                print(f"  Found {len(all_inputs)} input fields on page")
                
                for i, inp in enumerate(all_inputs):
                    if inp.is_displayed():
                        input_type = inp.get_attribute("type") or ""
                        input_id = inp.get_attribute("id") or ""
                        input_name = inp.get_attribute("name") or ""
                        input_placeholder = inp.get_attribute("placeholder") or ""
                        input_class = inp.get_attribute("class") or ""
                        
                        print(f"    Input {i}: type={input_type}, id={input_id}, name={input_name}, placeholder={input_placeholder}, class={input_class}")
                        
                        # Identify username field
                        if not username_field:
                            if (input_type in ["text", 'email', 'username'] or
                                "username" in input_id.lower() or "email" in input_id.lower() or
                                "username" in input_name.lower() or "email" in input_name.lower() or
                                "username" in input_placeholder.lower() or "email" in input_placeholder.lower() or
                                "username" in input_class.lower() or "email" in input_class.lower()):
                                username_field = inp
                                print(f"    Selected as username field: Input {i}")
                        
                        # Identify password field
                        if not password_field:
                            if (input_type == "password" or
                                "password" in input_id.lower() or
                                "password" in input_name.lower() or
                                "password" in input_placeholder.lower() or
                                "password" in input_class.lower()):
                                password_field = inp
                                print(f"    Selected as password field: Input {i}")
            
            # If we found both fields, proceed with login
            if username_field and password_field:
                # Clear and enter username
                try:
                    username_field.clear()
                    username_field.send_keys(credentials['username'])
                    print("  ✓ Username entered")
                except Exception as e:
                    print(f"  ❌ Error entering username: {e}")
                    continue
                
                # Clear and enter password
                try:
                    password_field.clear()
                    password_field.send_keys(credentials['password'])
                    print("  ✓ Password entered")
                except Exception as e:
                    print(f"  ❌ Error entering password: {e}")
                    continue
                
                # Find and click login button
                login_buttons = [
                    "//button[contains(., 'Login') or contains(., 'Sign In') or contains(., 'Log In') or contains(., 'Submit') or contains(., 'Continue')]",
                    "//input[@type='submit' and (contains(@value, 'Login') or contains(@value, 'Sign In') or contains(@value, 'Log In') or contains(@value, 'Submit'))]",
                    "//button[@type='submit']",
                    "//input[@type='submit']",
                    "//button[contains(@class, 'login') or contains(@class, 'submit') or contains(@class, 'continue')]",
                    "//input[contains(@class, 'login') or contains(@class, 'submit')]",
                    "//button[@data-id='login']",
                    "//span[contains(., 'Login') or contains(., 'Sign In')]/parent::button"
                ]
                
                login_button = None
                for button_selector in login_buttons:
                    try:
                        login_button = WebDriverWait(driver, 3).until(
                            EC.element_to_be_clickable((By.XPATH, button_selector)))
                        if login_button and login_button.is_displayed():
                            print(f"  Found login button with selector: {button_selector}")
                            break
                    except:
                        continue
                
                if login_button:
                    try:
                        login_button.click()
                        print("  ✓ Login button clicked")
                    except Exception as e:
                        print(f"  ❌ Error clicking login button: {e}")
                        # Try JavaScript click as fallback
                        try:
                            driver.execute_script("arguments[0].click();", login_button)
                            print("  ✓ Login button clicked via JavaScript")
                        except:
                            print("  ❌ Failed to click login button")
                else:
                    print("  No login button found, trying Enter key...")
                    password_field.send_keys(Keys.RETURN)
               
                # Wait for login to complete - with more flexible conditions
                try:
                    WebDriverWait(driver, 30).until(
                        lambda d: any(keyword in d.current_url.lower() for keyword in 
                                    ['dashboard', 'portal', 'solicitation', 'welcome', 'home', 'main', 'response', 'itsac']) or
                                any(keyword in d.page_source.lower() for keyword in 
                                    ['welcome', 'dashboard', 'solicitation', 'document', 'response', 'reference number', 'working title']),
                        message="Failed to reach portal dashboard page")
                   
                    print(f"  ✓ Login successful to portal!")
                    
                    # WAIT 25 SECONDS AFTER LOGIN BEFORE CONTINUING
                    print("  ⏳ Waiting 25 seconds after login for page to fully load...")
                    time.sleep(25)
                    print("  ✓ 25-second wait completed, continuing with content extraction...")
                    
                    return True
                    
                except Exception as e:
                    print(f"  ❌ Login verification failed: {str(e)}")
                    # Check if we're actually logged in by looking for specific elements
                    if any(keyword in driver.page_source.lower() for keyword in ['solicitation', 'response', 'reference number']):
                        print("  ✓ Found solicitation content, proceeding anyway...")
                        return True
                    else:
                        raise e
 
            else:
                print("  ❌ Could not find login fields after multiple attempts")
                return False
 
        except Exception as e:
            print(f"  ❌ Login attempt {attempt + 1}/{max_retries} failed: {str(e)}")
            if attempt < max_retries - 1:
                print("  Retrying in 5 seconds...")
                time.sleep(5)
                driver.get(url)
            else:
                print(f"  ✗ All login attempts failed for portal")
                return False
    pass

def extract_solicitation_response_number(driver):
    """Extract Solicitation Response number from the Salesforce Community page"""
    print("  🔍 Searching for Solicitation Response number...")
    
    try:
        # Wait for the main content to load
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.TAG_NAME, 'body')))
        
        # Get the page source
        page_html = driver.page_source
        soup = BeautifulSoup(page_html, 'html.parser')
        
        # Look for the specific Salesforce Community structure
        # The response number is in: lightning-formatted-text with value "230903DS"
        response_number = None
        
        # Method 1: Look for lightning-formatted-text element
        lightning_text = soup.find('lightning-formatted-text')
        if lightning_text and lightning_text.text.strip():
            response_number = lightning_text.text.strip()
            print(f"  ✓ Found response number in lightning-formatted-text: {response_number}")
        
        # Method 2: Look for the specific heading structure
        if not response_number:
            highlights_panel = soup.find('records-highlights2')
            if highlights_panel:
                primary_field = highlights_panel.find('lightning-formatted-text')
                if primary_field and primary_field.text.strip():
                    response_number = primary_field.text.strip()
                    print(f"  ✓ Found response number in highlights panel: {response_number}")
        
        # Method 3: Look for "Solicitation Reference Number" in tables
        if not response_number:
            tables = soup.find_all('table')
            for table in tables:
                table_text = table.get_text()
                if 'Solicitation Reference Number' in table_text:
                    # Extract the number after the label
                    match = re.search(r'Solicitation Reference Number:\s*([A-Z0-9]+)', table_text)
                    if match:
                        response_number = match.group(1)
                        print(f"  ✓ Found response number in table: {response_number}")
                        break
        
        # Method 4: Search the entire page for common patterns
        if not response_number:
            page_text = soup.get_text()
            patterns = [
                r'Solicitation Reference Number:\s*([A-Z0-9]+)',
                r'Reference Number:\s*([A-Z0-9]+)',
                r'Solicitation Number:\s*([A-Z0-9]+)',
                r'\b([A-Z]{2,}\d+[A-Z]*)\b'  # General pattern like "AB123CD"
            ]
            
            for pattern in patterns:
                matches = re.findall(pattern, page_text)
                if matches:
                    response_number = matches[0]
                    print(f"  ✓ Found response number with pattern {pattern}: {response_number}")
                    break
        
        if response_number:
            return response_number
        else:
            print("  ⚠️ Could not find solicitation response number, using timestamp")
            return f"UNKNOWN-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            
    except Exception as e:
        print(f"  ❌ Error extracting solicitation number: {str(e)}")
        return f"ERROR-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
    pass

def extract_salesforce_community_content(driver):
    """Extract complete content from Salesforce Community page"""
    try:
        print("  📄 Extracting Salesforce Community page content...")
        
        # Wait for page to be fully loaded
        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.TAG_NAME, 'body')))
        
        # Extract solicitation response number first
        solicitation_number = extract_solicitation_response_number(driver)
        
        # Scroll to load all content
        print("  📜 Scrolling page to load all content...")
        last_height = driver.execute_script("return document.body.scrollHeight")
        scroll_attempts = 0
        while scroll_attempts < 5:
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)
            new_height = driver.execute_script("return document.body.scrollHeight")
            if new_height == last_height:
                break
            last_height = new_height
            scroll_attempts += 1
        
        # Get complete page source
        page_html = driver.page_source
        soup = BeautifulSoup(page_html, 'html.parser')
        
        # Extract specific sections from the Salesforce Community structure
        
        # 1. Page Title and Header
        content = []
        content.append("=" * 80)
        content.append("SOLICITATION RESPONSE EXTRACTION - SALESFORCE COMMUNITY")
        content.append("=" * 80)
        content.append(f"Response Number: {solicitation_number}")
        
        # Extract page title
        title = soup.find('title')
        if title:
            content.append(f"Page Title: {title.get_text(strip=True)}")
        
        content.append(f"URL: {driver.current_url}")
        content.append(f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("=" * 80)
        content.append("")
        
        # 2. Main Content from the c-itsac-preview component
        preview_component = soup.find('c-itsac-preview')
        if preview_component:
            content.append("MAIN SOLICITATION CONTENT:")
            content.append("-" * 40)
            
            # Extract all text from the preview component
            preview_text = preview_component.get_text(separator='\n', strip=True)
            # Clean up the text
            lines = [line.strip() for line in preview_text.split('\n') if line.strip()]
            content.extend(lines)
            content.append("")
        
        # 3. Extract tables specifically
        tables = soup.find_all('table')
        if tables:
            content.append("TABLES FOUND:")
            content.append("-" * 40)
            for i, table in enumerate(tables, 1):
                content.append(f"Table {i}:")
                rows = table.find_all('tr')
                for row in rows:
                    cells = row.find_all(['td', 'th'])
                    row_data = [cell.get_text(strip=True) for cell in cells if cell.get_text(strip=True)]
                    if row_data:
                        content.append(" | ".join(row_data))
                content.append("")
        
        # 4. Extract key sections using specific class names or patterns
        key_sections = [
            ('HEADINGS', ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']),
            ('STRONG TEXT', ['strong']),
            ('LINKS', ['a']),
        ]
        
        for section_name, tags in key_sections:
            elements = soup.find_all(tags)
            if elements:
                content.append(f"{section_name}:")
                content.append("-" * 40)
                for elem in elements[:20]:  # Limit to first 20 of each type
                    text = elem.get_text(strip=True)
                    if text and len(text) > 3:
                        if section_name == 'LINKS' and elem.get('href'):
                            content.append(f"{text} -> {elem['href']}")
                        else:
                            content.append(text)
                content.append("")
        
        # 5. Extract the main content area
        main_content = soup.find('div', role='main')
        if main_content:
            main_text = main_content.get_text(separator='\n', strip=True)
            if main_text and len(main_text) > 100:
                content.append("MAIN CONTENT AREA:")
                content.append("-" * 40)
                lines = [line.strip() for line in main_text.split('\n') if line.strip() and len(line.strip()) > 10]
                content.extend(lines[:50])  # Limit to first 50 lines
                content.append("")
        
        # 6. If we still don't have much content, extract general text
        if len(content) < 20:
            content.append("GENERAL PAGE TEXT:")
            content.append("-" * 40)
            all_text = soup.get_text()
            lines = [line.strip() for line in all_text.split('\n') if line.strip()]
            # Filter out very short lines and navigation/text
            meaningful_lines = [line for line in lines if len(line) > 20 and not any(word in line.lower() for word in ['cookie', 'privacy', 'terms', 'login', 'sign'])]
            content.extend(meaningful_lines[:100])  # Limit to first 100 meaningful lines
        
        content.append("")
        content.append("=" * 80)
        content.append("END OF EXTRACTION")
        content.append("=" * 80)
        
        final_content = '\n'.join(content)
        print(f"  ✓ Extracted {len(final_content)} characters of content")
        
        return final_content, solicitation_number
        
    except Exception as e:
        print(f"  ❌ Error extracting page content: {str(e)}")
        error_content = f"Error extracting content: {str(e)}\n\nURL: {driver.current_url}\nPage Source Length: {len(driver.page_source)}"
        return error_content, f"ERROR-{datetime.now().strftime('%Y%m%d-%H%M%S')}"

    pass

def search_specific_hhsc_email(gmail_service, days_back=1):
    """Search for multiple department solicitation emails - ONLY TODAY'S EMAILS"""
    print(f"\n=== Searching for TODAY'S Department Solicitation Emails ===")
    
    today_date = datetime.now().strftime("%Y-%m-%d")
    print(f"📅 Will filter for emails from: {today_date}")
    
    # Use your EXACT department queries WITHOUT date filters first
    department_queries = [
        {
            'name': 'HHSC',
            'query': '"Texas Health and Human Services Commission has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'DFPS',
            'query': '"Texas Department of Family and Protective Services has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'TEA', 
            'query': '"Texas Education Agency has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'RRC',
            'query': '"Railroad Commission of Texas has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'OAG',
            'query': '"Office of the Attorney General of Texas has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'TDI',
            'query': '"Texas Department of Insurance has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'DPS',
            'query': '"Texas Department of Public Safety has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'OCCC',
            'query': '"Office of Consumer Credit Commissioner has posted Solicitation Number" "to the DIR ITSAC Portal"'
        },
        {
            'name': 'TSLAC',
            'query': '"Texas State Library and Archives Commission has posted Solicitation Number" "to the DIR ITSAC Portal"'
        }
    ]
    
    all_messages = []
    
    # REDUCED OUTPUT: Don't show individual department searches
    for dept in department_queries:
        try:
            results = gmail_service.users().messages().list(
                userId='me',
                q=dept['query'],
                maxResults=20
            ).execute()
            
            messages = results.get('messages', [])
            
            # Add department info to each message
            for message in messages:
                message['department'] = dept['name']
                message['query'] = dept['query']
            
            all_messages.extend(messages)
            
        except Exception as e:
            # Silent fail - just continue
            continue
    
    # REDUCED OUTPUT: Don't show total count
    # NOW FILTER FOR TODAY'S EMAILS ONLY
    today_messages = []
    
    # REDUCED OUTPUT: Don't show filtering progress
    for message in all_messages:
        try:
            # Get the email's actual date
            msg_data = gmail_service.users().messages().get(
                userId='me',
                id=message['id'],
                format='metadata',
                metadataHeaders=['Date']
            ).execute()
            
            internal_date = int(msg_data['internalDate'])
            email_date = datetime.fromtimestamp(internal_date / 1000)
            email_date_str = email_date.strftime("%Y-%m-%d")
            is_today = email_date.date() == datetime.now().date()
            
            if is_today:
                today_messages.append(message)
                # REDUCED: Don't show individual today's emails
            else:
                # REDUCED: Don't show individual old emails
                pass
                
        except Exception as e:
            # If we can't verify the date, include it to be safe
            today_messages.append(message)
    
    # SHOW ONLY FINAL RESULT
    if today_messages:
        print(f"🎯 Found {len(today_messages)} TODAY'S EMAILS TO PROCESS")
    else:
        print("📭 No TODAY'S emails found")
    
    return today_messages

def get_email_full_content(gmail_service, message_id):
    """Get complete email content including body and attachments"""
    try:
        message = gmail_service.users().messages().get(
            userId='me',
            id=message_id,
            format='full'
        ).execute()
        
        headers = message['payload'].get('headers', [])
        subject = next((header['value'] for header in headers if header['name'] == 'Subject'), 'No Subject')
        sender = next((header['value'] for header in headers if header['name'] == 'From'), 'Unknown Sender')
        date = next((header['value'] for header in headers if header['name'] == 'Date'), 'Unknown Date')
        
        full_body = extract_complete_email_body(message['payload'])
        
        return {
            'id': message_id,
            'subject': subject,
            'sender': sender,
            'date': date,
            'full_body': full_body
        }
        
    except Exception as e:
        print(f"Error getting email details: {e}")
        return None
    pass

def extract_complete_email_body(payload):
    """Extract complete text content from email payload including HTML"""
    body = ""
    
    def process_part(part):
        nonlocal body
        if part['mimeType'] == 'text/plain':
            data = part['body'].get('data', '')
            if data:
                body += base64.urlsafe_b64decode(data).decode('utf-8') + "\n"
        elif part['mimeType'] == 'text/html':
            data = part['body'].get('data', '')
            if data:
                html_content = base64.urlsafe_b64decode(data).decode('utf-8')
                soup = BeautifulSoup(html_content, 'html.parser')
                for element in soup.find_all(['p', 'br', 'a']):
                    if element.name == 'br':
                        body += "\n"
                    elif element.name == 'a' and element.get('href'):
                        body += f"{element.get_text()} [LINK: {element['href']}]\n"
                    else:
                        body += element.get_text() + "\n"
    
    if 'parts' in payload:
        for part in payload['parts']:
            process_part(part)
            if 'parts' in part:
                for subpart in part['parts']:
                    process_part(subpart)
    else:
        process_part(payload)
    
    return body.strip()

    pass

def extract_portal_link(email_body):
    """Extract the portal link from the email body"""
    print("\n=== Searching for Portal Link ===")
    
    patterns = [
        r'If you would like to respond please go to:\s*([^\s]+)',
        r'please go to:\s*([^\s]+)',
        r'portal link:\s*([^\s]+)',
        r'DIR ITSAC Portal[^:]*:\s*([^\s]+)',
        r'https?://[^\s]+'
    ]
    
    for pattern in patterns:
        matches = re.findall(pattern, email_body, re.IGNORECASE | re.MULTILINE)
        if matches:
            print(f"Found link using pattern: {pattern}")
            for match in matches:
                url = match.strip()
                if url.endswith(('.', ',', ';', ':')):
                    url = url[:-1]
                if url.startswith(('http://', 'https://')):
                    print(f"✅ Valid portal link found: {url}")
                    return url
    
    print("Searching for any URLs in email body...")
    url_pattern = r'https?://[^\s]+'
    all_urls = re.findall(url_pattern, email_body)
    
    for url in all_urls:
        clean_url = url.strip()
        if clean_url.endswith(('.', ',', ';', ':')):
            clean_url = clean_url[:-1]
        print(f"Found URL: {clean_url}")
        if any(keyword in clean_url.lower() for keyword in ['dir', 'itsac', 'portal', 'solicitation', 'hhsc']):
            print(f"✅ Portal-like link found: {clean_url}")
            return clean_url
    
    print("❌ No portal link found in email body")
    return None

    pass

def display_email_info(email_details, department):
    """Display email information with department"""
    print(f"\n{'='*100}")
    print(f"📧 {department} EMAIL FOUND!")
    print(f"{'='*100}")
    print(f"📋 SUBJECT: {email_details['subject']}")
    print(f"👤 FROM: {email_details['sender']}")
    print(f"📅 DATE: {email_details['date']}")
    print(f"🏢 DEPARTMENT: {department}")
    print(f"{'='*100}")
    pass

def logout_from_portal(driver):
    """Logout from the current portal session"""
    print("  🚪 Attempting to logout from current portal...")
    try:
        # Try multiple logout methods
        logout_selectors = [
            "//a[contains(., 'Logout') or contains(., 'Log Out') or contains(., 'Sign Out') or contains(., 'Signout')]",
            "//button[contains(., 'Logout') or contains(., 'Log Out') or contains(., 'Sign Out') or contains(., 'Signout')]",
            "//span[contains(., 'Logout') or contains(., 'Log Out')]/parent::button",
            "//div[contains(@class, 'logout') or contains(@class, 'signout')]",
            "//a[@href*='logout' or @href*='Logout']",
            "//*[@data-id='logout' or contains(@onclick, 'logout')]"
        ]
        
        for selector in logout_selectors:
            try:
                logout_element = WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, selector)))
                if logout_element and logout_element.is_displayed():
                    logout_element.click()
                    print("  ✓ Logout button clicked")
                    
                    # Wait for logout to complete (redirect to login page)
                    WebDriverWait(driver, 10).until(
                        lambda d: any(keyword in d.current_url.lower() for keyword in 
                                    ['login', 'signin', 'auth', 'salesforce']) or
                                any(keyword in d.page_source.lower() for keyword in 
                                    ['login', 'sign in', 'username', 'password']))
                    print("  ✓ Logout successful")
                    time.sleep(3)  # Additional wait after logout
                    return True
            except:
                continue
        
        # If no logout button found, try to clear session by going to logout URL
        try:
            driver.get("https://dir.my.site.com/itsacvendor/secur/logout.jsp")
            print("  ✓ Manual logout via URL")
            time.sleep(3)
            return True
        except:
            pass
        
        print("  ⚠️ Could not find logout button, clearing cookies instead")
        driver.delete_all_cookies()
        driver.refresh()
        time.sleep(3)
        return True
        
    except Exception as e:
        print(f"  ❌ Error during logout: {str(e)}")
        # Fallback: clear cookies and refresh
        try:
            driver.delete_all_cookies()
            driver.refresh()
            print("  ✓ Cleared cookies as fallback")
            time.sleep(3)
            return True
        except:
            print("  ❌ Failed to clear session")
            return False
    pass

def process_portal_without_login(driver, portal_link, department="HHSC"):
    """Process portal content without logging in again (assumes already logged in)"""
    print(f"\n=== PROCESSING {department} PORTAL CONTENT ===")
    print(f"Portal URL: {portal_link}")
    
    try:
        # Navigate to the portal link (should maintain session)
        driver.get(portal_link)
        time.sleep(5)
        
        # Wait for page to load
        WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.TAG_NAME, 'body')))
        
        # Extract Salesforce Community content
        portal_content, solicitation_number = extract_salesforce_community_content(driver)
        
        output_dir = "dir_portal_outputs"
        # Clear and create output directory (if needed)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # Clean the solicitation number for filename use
        clean_solicitation_number = re.sub(r'[^a-zA-Z0-9-]', '_', solicitation_number)
        
        # Use format: "Solicitation_Response_Number_529601512"
        output_file = os.path.join(output_dir, f"Solicitation_Response_Number_{clean_solicitation_number}.txt")
        
        # Add department info to the content
        portal_content_with_dept = f"DEPARTMENT: {department}\n" + portal_content
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(portal_content_with_dept)
        
        print(f"✅ {department} portal content saved to: {output_file}")
        print(f"✅ Solicitation Response Number: {solicitation_number}")

        # Download forms after extracting content
        print(f"\n📄 Attempting to download forms for {solicitation_number}...")
        forms_downloaded = download_forms_from_portal(driver, solicitation_number, department)
        if forms_downloaded:
            print(f"✅ Forms downloaded successfully for {solicitation_number}")
        else:
            print(f"⚠️  Could not download forms for {solicitation_number}")

        print(f"✅ {department} portal processing completed successfully!")
        
        return True
        
    except Exception as e:
        print(f"❌ Error processing {department} portal content: {str(e)}")
        return False

    pass

def download_forms_from_portal(driver, solicitation_number, department):
    """Click Download Forms button and wait for complete download"""
    print(f"  📥 Attempting to download forms for {department}...")
    
    try:
        # First, find and click the "Download Forms" TAB
        download_tab_selectors = [
            "//a[contains(., 'Download Forms')]",
            "//button[contains(., 'Download Forms')]",
            "//span[contains(., 'Download Forms')]/parent::button",
            "//li[contains(., 'Download Forms')]",
            "//*[contains(text(), 'Download Forms') and contains(@class, 'tab')]",
        ]
        
        download_tab_clicked = False
        for selector in download_tab_selectors:
            try:
                download_tab = WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, selector)))
                if download_tab and download_tab.is_displayed():
                    print(f"  Found Download Forms tab with selector: {selector}")
                    driver.execute_script("arguments[0].click();", download_tab)
                    print("  ✓ Clicked Download Forms tab")
                    time.sleep(3)
                    download_tab_clicked = True
                    break
            except:
                continue
        
        if not download_tab_clicked:
            print("  ⚠️ Could not find Download Forms tab")
            return False
        
        # Now look for the actual download button
        print("  🔍 Looking for download button...")
        
        download_button_selectors = [
            "//button[contains(., 'Download Forms') and not(contains(@class, 'tab'))]",
            "//a[contains(., 'Download Forms')]",
            "//button[contains(., 'Download') and contains(@class, 'download')]",
            "//lightning-button[contains(., 'Download Forms')]//button",
        ]
        
        for selector in download_button_selectors:
            try:
                download_button = WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, selector)))
                if download_button and download_button.is_displayed():
                    print(f"  Found download button with selector: {selector}")
                    
                    # Click the download button
                    driver.execute_script("arguments[0].click();", download_button)
                    print("  ✓ Clicked Download Forms button")
                    
                    # WAIT FOR DOWNLOAD TO COMPLETE (enhanced waiting)
                    documents_dir = os.path.join(os.getcwd(), "dir_documents")
                    wait_for_download_completion(documents_dir, timeout=30)
                    
                    # Check if file was downloaded successfully
                    success = check_downloaded_file_quick(solicitation_number, department)
                    
                    if success:
                        print("  ✅ Download and verification completed")
                        # Process the document immediately after download
                        print("  🔧 Processing document content...")
                        process_downloaded_document(solicitation_number)
                    else:
                        print("  ⚠️ Download verification failed")
                    
                    return success
                    
            except Exception as e:
                continue
        
        print("  ❌ Could not find downloadable forms button")
        return False
        
    except Exception as e:
        print(f"  ❌ Error downloading forms: {str(e)}")
        return False
    pass

def wait_for_download_completion(documents_dir, timeout=30):
    """Wait for all downloads to complete in the directory"""
    print("  ⏳ Waiting for all downloads to complete...")
    
    start_time = time.time()
    while time.time() - start_time < timeout:
        # Check for .crdownload files (Chrome download in progress)
        download_in_progress = False
        for file in os.listdir(documents_dir):
            if file.endswith('.crdownload') or file.endswith('.tmp'):
                download_in_progress = True
                break
        
        if not download_in_progress:
            # Check if any document files are complete
            doc_files = [f for f in os.listdir(documents_dir) 
                        if f.lower().endswith(('.doc', '.docx', '.pdf'))]
            if doc_files:
                complete_files = [f for f in doc_files 
                                if is_file_complete(os.path.join(documents_dir, f))]
                if complete_files:
                    print("  ✅ All downloads completed")
                    return True
        
        time.sleep(2)
    
    print("  ⚠️ Download wait timeout - some files may be incomplete")
    return False

    pass

def is_file_complete(file_path):
    """Check if file is completely downloaded (not being written to)"""
    try:
        # Check if file exists and has size > 0
        if not os.path.exists(file_path) or os.path.getsize(file_path) == 0:
            return False
        
        # Try to open the file in exclusive mode - if we can, it's not being written to
        try:
            with open(file_path, 'a+b') as f:
                # Try to get an exclusive lock
                import msvcrt
                msvcrt.locking(f.fileno(), msvcrt.LK_NBLCK, 1)
                msvcrt.locking(f.fileno(), msvcrt.LK_UNLCK, 1)
                return True
        except (IOError, OSError):
            # File is locked by another process (still downloading)
            return False
        
    except Exception:
        return False

    pass

def check_downloaded_file_quick(solicitation_number, department):
    """Enhanced download verification with file completion checks"""
    try:
        documents_dir = os.path.join(os.getcwd(), "dir_documents")
        
        # Wait for downloads to complete with timeout
        max_wait_time = 30  # Maximum wait time in seconds
        wait_interval = 2   # Check every 2 seconds
        total_waited = 0
        
        print(f"  ⏳ Waiting for download to complete...")
        
        while total_waited < max_wait_time:
            # Look for any .doc, .docx, or .pdf files
            found_files = []
            for file in os.listdir(documents_dir):
                if file.lower().endswith(('.doc', '.docx', '.pdf')):
                    file_path = os.path.join(documents_dir, file)
                    
                    # Check if file is complete (not being downloaded)
                    if is_file_complete(file_path):
                        found_files.append(file_path)
                    else:
                        print(f"  ⏳ File still downloading: {file}")
            
            if found_files:
                # Use the most recently modified file
                latest_file = max(found_files, key=os.path.getmtime)
                
                # Clean filename and rename to match solicitation number
                clean_solicitation_number = re.sub(r'[^a-zA-Z0-9-]', '_', solicitation_number)
                
                # Get file extension
                file_ext = os.path.splitext(latest_file)[1].lower()
                new_filename = f"Solicitation_Response_Number_{clean_solicitation_number}_Forms{file_ext}"
                new_filepath = os.path.join(documents_dir, new_filename)
                
                # Rename the file
                if latest_file != new_filepath:
                    os.rename(latest_file, new_filepath)
                
                print(f"  ✅ Download completed: {new_filename}")
                return True
            
            # Wait before checking again
            time.sleep(wait_interval)
            total_waited += wait_interval
            print(f"  ⏳ Still waiting... ({total_waited}s/{max_wait_time}s)")
        
        print("  ❌ Download timeout - file may be incomplete")
        return False
        
    except Exception as e:
        print(f"  ❌ Error in file check: {str(e)}")
        return False

    pass

def process_downloaded_document(solicitation_number):
    """Open downloaded document and remove everything before 'SOLICITATION CONTACT' for all file types"""
    try:
        documents_dir = "hhsc_documents"
        clean_solicitation_number = re.sub(r'[^a-zA-Z0-9-]', '_', solicitation_number)
        
        processed_files = []
        
        # Look for all document files with this solicitation number
        for file in os.listdir(documents_dir):
            if file.lower().endswith(('.docx', '.doc', '.pdf')) and clean_solicitation_number in file:
                filepath = os.path.join(documents_dir, file)
                
                if file.lower().endswith('.docx'):
                    if process_docx_file(filepath):
                        processed_files.append(file)
                        print(f"  ✅ Processed DOCX: {file}")
                
                elif file.lower().endswith('.doc'):
                    if process_doc_file(filepath):
                        processed_files.append(file)
                        print(f"  ✅ Processed DOC: {file}")
                
                elif file.lower().endswith('.pdf'):
                    if process_pdf_file(filepath):
                        processed_files.append(file)
                        print(f"  ✅ Processed PDF: {file}")
        
        if processed_files:
            print(f"  ✅ Successfully processed {len(processed_files)} files for {solicitation_number}")
            return True
        else:
            print(f"  ⚠️ No files found to process for {solicitation_number}")
            return False
        
    except Exception as e:
        print(f"  ❌ Error processing document: {str(e)}")
        return False


def process_docx_file(filepath):
    """Remove everything from start up to and including SOLICITATION CONTACT section, remove header tables, clean up blank pages, and add proper page breaks between sections"""
    try:
        doc = Document(filepath)
        
        # Find the paragraph containing "XII. SOLICITATION CONTACT"
        solicitation_contact_index = -1
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.upper().strip()
            
            # Find the "XII. SOLICITATION CONTACT" section
            if "XII." in text and "SOLICITATION CONTACT" in text:
                solicitation_contact_index = i
                print(f"  📝 Found 'XII. SOLICITATION CONTACT' at position {i}")
                
                # Find where this section ends (typically 2-4 paragraphs after)
                section_end_index = i
                for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
                    next_text = doc.paragraphs[j].text.strip()
                    if not next_text:
                        continue
                    # Look for the actual contact info and stop when we find empty lines or next section
                    if "@" in next_text or "DIR WEB SITE" in next_text.upper() or "DIR WEBSITE" in next_text.upper():
                        section_end_index = j
                    # Stop when we find candidate sections or major headings
                    if "CANDIDATE" in next_text.upper() or re.match(r'^[XIV]+\.', next_text.upper()):
                        break
                
                # Make sure we include the contact information
                if section_end_index == i:
                    section_end_index = i + 4  # Include contact info paragraphs               
                break
        
        if solicitation_contact_index == -1:
            # Try alternative pattern without "XII."
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.upper().strip()
                if "SOLICITATION CONTACT" in text:
                    solicitation_contact_index = i
                    print(f"  📝 Found 'SOLICITATION CONTACT' at position {i}")
                    
                    # Find where this section ends
                    section_end_index = i
                    for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
                        next_text = doc.paragraphs[j].text.strip()
                        if not next_text:
                            continue
                        if "@" in next_text or "DIR WEB SITE" in next_text.upper():
                            section_end_index = j
                        if "CANDIDATE" in next_text.upper():
                            break
                    
                    if section_end_index == i:
                        section_end_index = i + 4                  
                    break
        
        if solicitation_contact_index == -1:
            print(f"  ⚠️ 'SOLICITATION CONTACT' section not found in {os.path.basename(filepath)}, keeping original")
            return True
        
        print(f"  📝 Removing everything from start to position {section_end_index}")
        
        # Remove all paragraphs from beginning up to and including SOLICITATION CONTACT section
        paragraphs_to_remove = list(range(0, section_end_index + 1))
        
        # Remove paragraphs in reverse order to avoid index issues
        for i in sorted(paragraphs_to_remove, reverse=True):
            if i < len(doc.paragraphs):
                p = doc.paragraphs[i]
                if p._element is not None:
                    p._element.getparent().remove(p._element)
        
        # Remove the header table and Minimum Requirements table
        tables_removed = 0
        for table in doc.tables[:]:  # Use slice copy to avoid modification during iteration
            table_text = ""
            # Get all text from the table
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
            
            table_text_upper = table_text.upper()
            
            # Check if this is the HEADER table (solicitation reference table)
            if ("SOLICITATION REFERENCE" in table_text_upper and 
                "WORKING TITLE" in table_text_upper and 
                "TITLE/LEVEL" in table_text_upper and
                "CATEGORY" in table_text_upper and
                "NTE RATE" in table_text_upper):
                
                # Remove this header table
                tbl = table._element
                tbl.getparent().remove(tbl)
                tables_removed += 1
                print(f"  ✅ Removed Header Table (Solicitation Reference table)")
                continue  # Continue to check for other tables to remove
            
            # Check if this is the Minimum Requirements table
            if ("MINIMUM REQUIREMENTS" in table_text_upper and 
                "YEARS" in table_text_upper and 
                "REQUIRED/PREFERRED" in table_text_upper and
                "EXPERIENCE" in table_text_upper):
                
                # Remove this table
                tbl = table._element
                tbl.getparent().remove(tbl)
                tables_removed += 1
                print(f"  ✅ Removed Minimum Requirements table")
        
        print(f"  ✅ Kept {len(doc.tables)} candidate-related tables, removed {tables_removed} tables total")
        
        # Remove ALL content before the first "CANDIDATE" section
        print(f"  🧹 Removing ALL content before first CANDIDATE section...")
        
        # Find the first occurrence of any CANDIDATE section
        first_candidate_index = -1
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.upper().strip()
            if "CANDIDATE" in text:
                first_candidate_index = i
                print(f"  📝 Found first CANDIDATE section at position {i}: {text}")
                break
        
        if first_candidate_index > 0:
            # Remove everything before the first CANDIDATE section
            paragraphs_to_remove_before_candidate = list(range(0, first_candidate_index))
            for i in sorted(paragraphs_to_remove_before_candidate, reverse=True):
                if i < len(doc.paragraphs):
                    p = doc.paragraphs[i]
                    if p._element is not None:
                        p._element.getparent().remove(p._element)
            print(f"  ✅ Removed {len(paragraphs_to_remove_before_candidate)} paragraphs before first CANDIDATE section")
        
        # NEW: Remove blank pages BETWEEN sections
        print(f"  🧹 Removing blank pages between sections...")
        blank_paragraphs_removed = 0
        
        # Remove empty paragraphs throughout the entire document
        i = 0
        while i < len(doc.paragraphs):
            paragraph_text = doc.paragraphs[i].text.strip()
            
            # Remove paragraphs that are empty or contain only special characters
            if (not paragraph_text or 
                paragraph_text in ['-', '–', '—', '•', '*', '·', ' '] or
                (len(paragraph_text) < 3 and not any(c.isalnum() for c in paragraph_text))):
                
                # Check if this empty paragraph is between major sections
                if i > 0 and i < len(doc.paragraphs) - 1:
                    prev_text = doc.paragraphs[i-1].text.upper().strip() if i > 0 else ""
                    next_text = doc.paragraphs[i+1].text.upper().strip() if i < len(doc.paragraphs) - 1 else ""
                    
                    # If this empty paragraph is between two candidate sections, remove it
                    if ("CANDIDATE" in prev_text and "CANDIDATE" in next_text):
                        p = doc.paragraphs[i]
                        if p._element is not None:
                            p._element.getparent().remove(p._element)
                            blank_paragraphs_removed += 1
                            print(f"  ✅ Removed blank page between sections at position {i}")
                            continue  # Don't increment i since we removed current element
                
                # Also remove standalone empty paragraphs
                p = doc.paragraphs[i]
                if p._element is not None:
                    p._element.getparent().remove(p._element)
                    blank_paragraphs_removed += 1
                    continue  # Don't increment i since we removed current element
            
            i += 1
        
        # Clean up any remaining empty paragraphs at the beginning
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            first_paragraph = doc.paragraphs[0]
            if first_paragraph._element is not None:
                first_paragraph._element.getparent().remove(first_paragraph._element)
        
        # NEW: Add proper page breaks between major sections
        print(f"  📄 Adding proper page breaks between major sections...")
        page_breaks_added = 0
        
        # Find all major section headings
        section_headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.upper().strip()
            if "CANDIDATE REFERENCE" in text:
                section_headings.append(("CANDIDATE REFERENCE", i))
            elif "CANDIDATE QUALIFICATIONS" in text:
                section_headings.append(("CANDIDATE QUALIFICATIONS", i))
            elif "CANDIDATE ACKNOWLEDGEMENT" in text:
                section_headings.append(("CANDIDATE ACKNOWLEDGEMENT", i))
        
        # Sort by position
        section_headings.sort(key=lambda x: x[1])
        
        # Add page breaks before each section (except the first one)
        if len(section_headings) > 1:
            for i in range(1, len(section_headings)):
                section_name, position = section_headings[i]
                
                # Use paragraph formatting for page breaks (more reliable)
                doc.paragraphs[position].paragraph_format.page_break_before = True
                page_breaks_added += 1
                print(f"  ✅ Added page break before {section_name}")
        
        # Final cleanup: ensure no blank pages at the beginning
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            first_paragraph = doc.paragraphs[0]
            if first_paragraph._element is not None:
                first_paragraph._element.getparent().remove(first_paragraph._element)
        
        # Save the modified document (overwrite original)
        doc.save(filepath)
        print(f"  ✅ Removed everything up to and including SOLICITATION CONTACT section")
        print(f"  ✅ Removed Header Table and Minimum Requirements table")
        print(f"  ✅ Removed {blank_paragraphs_removed} blank pages between sections")
        print(f"  ✅ Added {page_breaks_added} proper page breaks between sections")
        print(f"  ✅ Document now has clean 3-page structure")
        return True
        
    except Exception as e:
        print(f"  ❌ Error processing DOCX file {filepath}: {str(e)}")
        return False

    

def process_pdf_file(filepath):
    """Remove everything from start up to and including SOLICITATION CONTACT section in PDF"""
    try:
        # Read the PDF
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Find the page containing "XII. SOLICITATION CONTACT"
            solicitation_contact_page_index = -1
            pages_to_keep_start = -1
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                
                # Look for SOLICITATION CONTACT section
                if "XII." in text and "SOLICITATION CONTACT" in text:
                    solicitation_contact_page_index = page_num
                    print(f"  📝 Found 'XII. SOLICITATION CONTACT' on page {page_num + 1}")
                    
                    # Find where the contact section ends on this page
                    lines = text.split('\n')
                    for i, line in enumerate(lines):
                        if "CANDIDATE" in line.upper() and "REFERENCE" in line.upper():
                            # Found candidate reference section - keep from this page onward
                            pages_to_keep_start = page_num
                            break
                    
                    if pages_to_keep_start == -1:
                        # If we didn't find candidate reference on this page, check next page
                        pages_to_keep_start = page_num + 1
                    
                    break
                elif "SOLICITATION CONTACT" in text:
                    # Alternative pattern without "XII."
                    solicitation_contact_page_index = page_num
                    print(f"  📝 Found 'SOLICITATION CONTACT' on page {page_num + 1}")
                    
                    # Look for candidate sections
                    if "CANDIDATE" in text.upper():
                        pages_to_keep_start = page_num
                    else:
                        pages_to_keep_start = page_num + 1
                    break
            
            if solicitation_contact_page_index == -1:
                print(f"  ⚠️ 'SOLICITATION CONTACT' section not found in PDF {os.path.basename(filepath)}, keeping original")
                return True
            
            # Create a new PDF with pages starting after SOLICITATION CONTACT section
            pdf_writer = PyPDF2.PdfWriter()
            
            # Add all pages starting from the candidate sections
            for page_num in range(pages_to_keep_start, len(pdf_reader.pages)):
                pdf_writer.add_page(pdf_reader.pages[page_num])
            
            # Save the modified PDF (overwrite original)
            with open(filepath, 'wb') as output_file:
                pdf_writer.write(output_file)
            
            removed_pages = pages_to_keep_start
            kept_pages = len(pdf_reader.pages) - pages_to_keep_start
            print(f"  ✅ Removed {removed_pages} pages up to SOLICITATION CONTACT, kept {kept_pages} pages with candidate sections")
            return True
            
    except Exception as e:
        print(f"  ❌ Error processing PDF file {filepath}: {str(e)}")
        return False

    
def process_doc_file(filepath):
    """Remove everything from start up to and including SOLICITATION CONTACT section in DOC file"""
    try:
        # Convert DOC to DOCX first
        docx_filepath = filepath.replace('.doc', '.docx')
        
        print(f"  🔄 Converting DOC to DOCX: {os.path.basename(filepath)}")
        
        # Use win32com to convert .doc to .docx
        pythoncom.CoInitialize()
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        
        try:
            doc = word_app.Documents.Open(os.path.abspath(filepath))
            doc.SaveAs2(os.path.abspath(docx_filepath), FileFormat=16)  # 16 = wdFormatDocumentDefault
            doc.Close()
            print(f"  ✅ Conversion successful")
        finally:
            word_app.Quit()
            pythoncom.CoUninitialize()
        
        # Process the converted DOCX file
        if process_docx_file(docx_filepath):
            # Remove the original DOC file and keep the processed DOCX
            os.remove(filepath)
            print(f"  ✅ Removed original DOC file, kept processed DOCX")
            return True
        return False
        
    except Exception as e:
        print(f"  ❌ Error processing DOC file {filepath}: {str(e)}")
        return False

def process_all_downloaded_documents():
    """Process all downloaded documents after portal processing is complete"""
    print(f"\n📄 Processing all downloaded documents...")
    
    documents_dir = "dir_documents"
    if not os.path.exists(documents_dir):
        print("  ⚠️ No documents folder found")
        return
    
    # Get all document files
    doc_files = []
    for file in os.listdir(documents_dir):
        if file.lower().endswith(('.doc', '.docx', '.pdf')):
            doc_files.append(file)
    
    if not doc_files:
        print("  ⚠️ No documents found to process")
        return
    
    processed_count = 0
    
    # Process by solicitation number patterns
    for file in doc_files:
        # Extract solicitation number from filename
        match = re.search(r'Solicitation_Response_Number_([^_\.]+)', file)
        if match:
            solicitation_number = match.group(1)
            if process_downloaded_document(solicitation_number):
                processed_count += 1
    
    print(f"  ✅ Processed {processed_count} out of {len(doc_files)} documents")
    

def main():
    """Main function to search for department emails and automate portal access with intelligent monitoring"""
    try:
        print("🚀 Starting Intelligent Solicitation Automation System")
        print("=" * 70)
        print("🎯 WILL PROCESS ONLY NEWLY RECEIVED EMAILS")
        print("📧 Existing emails will be marked as processed")
        print("🗑️  Folders will be cleared when new emails are found")
        print("🔄 Continuous monitoring for new arrivals")
        print("=" * 70)
        
        # Clear output folders at the start
        output_dirs = ["dir_portal_outputs", "dir_documents"]
        
        for output_dir in output_dirs:
            if os.path.exists(output_dir):
                for file in os.listdir(output_dir):
                    file_path = os.path.join(output_dir, file)
                    try:
                        if os.path.isfile(file_path):
                            os.unlink(file_path)
                            print(f"🗑️  Deleted previous file: {file}")
                    except Exception as e:
                        print(f"❌ Error deleting {file_path}: {e}")
                print(f"✅ Cleared previous files in {output_dir}")
            else:
                os.makedirs(output_dir)
                print(f"✅ Created directory: {output_dir}")
        
        # Initialize automation FIRST to track processed emails
        automation = SolicitationAutomation()
        
        # Initialize driver and Gmail service for initial processing
        driver = initialize_driver()
        gmail_service = auto_authenticate_primary_gmail()
        
        # Search for department emails - ONLY TODAY'S
        today_date = datetime.now().strftime("%Y-%m-%d")
        print(f"📅 Searching for TODAY'S emails only: {today_date}")
        messages = search_specific_hhsc_email(gmail_service, days_back=1)
        
        initial_processed = 0
        
        if messages:
            print(f"🎯 Found {len(messages)} TODAY'S emails for initial processing")
            
            # ⭐⭐ MARK all existing emails as processed to avoid re-processing later
            for message in messages:
                automation.processed_emails.add(message['id'])
            print(f"📝 Marked {len(messages)} existing emails as processed to avoid duplicates")
            
            processed_portals = 0
            
            # NEW APPROACH: Login once, process all departments, logout once
            print(f"\n🚀 SINGLE LOGIN FOR ALL TODAY'S DEPARTMENTS")
            
            # Get the first portal link to login
            first_portal_link = None
            for i, message in enumerate(messages):
                email_details = get_email_full_content(gmail_service, message['id'])
                if email_details:
                    portal_link = extract_portal_link(email_details['full_body'])
                    if portal_link:
                        first_portal_link = portal_link
                        break
            
            if not first_portal_link:
                print("❌ No valid portal links found in any TODAY'S emails")
                driver.quit()
            else:
                # Login once at the beginning
                print(f"\n=== SINGLE LOGIN FOR ALL TODAY'S DEPARTMENTS ===")
                login_success = login_to_hhsc_portal(driver, first_portal_link, DEPARTMENT_CREDENTIALS)
                
                if not login_success:
                    print("❌ Failed to login to portal")
                    driver.quit()
                else:
                    print("✅ Successfully logged in! Now processing all TODAY'S departments...")
                    
                    # Process all departments with the same session
                    for i, message in enumerate(messages, 1):
                        print(f"\n📥 Processing TODAY'S email {i}/{len(messages)}...")
                        
                        email_details = get_email_full_content(gmail_service, message['id'])
                        department = message.get('department', 'Unknown Department')
                        
                        if not email_details:
                            print("❌ Failed to get email content")
                            continue
                            
                        display_email_info(email_details, department)
                        portal_link = extract_portal_link(email_details['full_body'])
                        
                        if portal_link:
                            print(f"🔗 Portal link extracted: {portal_link}")
                            print(f"🏢 Department: {department}")
                            
                            # Process the portal without logging in again (maintain session)
                            success = process_portal_without_login(driver, portal_link, department)
                            
                            if success:
                                processed_portals += 1
                                print(f"✅ Successfully processed {department} portal {processed_portals}")
                                
                                # Short delay between departments
                                if i < len(messages):
                                    print("⏳ Quick preparation for next department in 3 seconds...")
                                    time.sleep(3)
                            else:
                                print(f"❌ Failed to process {department} portal")
                        else:
                            print(f"❌ No portal link found in {department} email")
                    
                    # Logout once at the end
                    print(f"\n=== SINGLE LOGOUT AFTER ALL TODAY'S DEPARTMENTS ===")
                    logout_success = logout_from_portal(driver)
                    if logout_success:
                        print("✅ Successfully logged out after processing all TODAY'S departments")
                    else:
                        print("⚠️  Logout had issues, but processing completed")
                    
                    print(f"\n🎉 TODAY'S PORTAL PROCESSING COMPLETED!")
                    print(f"   Found {len(messages)} TODAY'S department email(s)")
                    print(f"   Automatically processed {processed_portals} portal(s)")
                    
                    initial_processed = processed_portals
        else:
            print("❌ No TODAY'S department emails found for initial processing.")
            driver.quit()
        
        # Process all downloaded documents at the end
        if initial_processed > 0:
            process_all_downloaded_documents()
            
            # REGEX EXTRACTION - ADDED BEFORE LLM PROCESSING
            print(f"\n=== REGEX DATA EXTRACTION ===")
            regex_processed = extract_data_with_regex()
            print(f"✅ Regex extraction completed: {regex_processed} files processed")
            
            # ADD LLM PROCESSING AFTER ALL PORTAL PROCESSING
            print(f"\n=== STARTING LLM DATA PROCESSING ===")
            process_files_with_llm()

            # EMAIL SENDING
            print(f"\n=== SENDING SOLICITATION RESPONSE EMAILS ===")
            email_sender = HHSCOutlookEmailSender()
            emails_sent = email_sender.send_emails_for_all_solicitations()
            print(f"✅ Sent {emails_sent} solicitation response emails")
            
            print(f"\n✅ INITIAL PROCESSING COMPLETED: {initial_processed} solicitations processed")
        else:
            print("\n📭 No initial solicitations to process")
        
        # ===== START INTELLIGENT CONTINUOUS AUTOMATION =====
        print(f"\n{'='*70}")
        print("🤖 STARTING INTELLIGENT CONTINUOUS AUTOMATION")
        print(f"{'='*70}")
        print("📧 Monitoring for NEW emails (not previously processed)")
        print("⏰ Checking every 5 minutes")
        print("🎯 Will ONLY process NEWLY RECEIVED emails")
        print("🗑️  Folders will be CLEARED when new emails are found")
        print("🔄 Automation will run CONTINUOUSLY")
        print("📅 Today's date: " + datetime.now().strftime("%Y-%m-%d"))
        print("⏹️  Press Ctrl+C to stop automation")
        print(f"{'='*70}\n")
        
        # Start intelligent continuous monitoring with the SAME automation instance
        automation.start_automation()
        
        try:
            # Keep running with status updates - INDEFINITELY
            while automation.is_running:
                time.sleep(30)  # Check status every 30 seconds
                
                # Show status every 5 minutes
                if int(time.time()) % 300 == 0:
                    status = automation.get_status()
                    print(f"\n📊 Automation Status Update:")
                    print(f"   ✅ Processed today: {status['processed_count']} emails")
                    print(f"   🔄 Status: {status['status']}")
                    print(f"   ⏰ Last check: {status['last_check']}")
                    print(f"   📧 Still monitoring for NEW incoming emails...")
                    
        except KeyboardInterrupt:
            print("\n🛑 User requested to stop automation...")
            automation.stop_automation()
            print("✅ Automation stopped successfully")
            print("🏁 Script execution finished.")
            
    except Exception as e:
        print(f"❌ Error in main function: {e}")
        import traceback
        traceback.print_exc()
    

if __name__ == "__main__":
    main()